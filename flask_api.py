from flask import Flask, request, jsonify, send_file 
import pandas as pd
import io
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import tempfile
import os
from openai import OpenAI
from typing import Dict, List, Optional
import plotly.express as px
import plotly.graph_objects as go
from PIL import Image
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.io as pio
import numpy as np
import json
from werkzeug.utils import secure_filename
import logging
import threading
import time
from datetime import datetime
import requests
import urllib.parse
import tempfile
import os
from bs4 import BeautifulSoup
from flask_cors import CORS


# Try to import credentials, fall back to environment variables if not available
try:
    from credentials import DOMAIN, CLIENT_ID, CLIENT_SECRET, USERNAME, PASSWORD, ROOT_FOLDER, OPENAI_API_KEY
    EGNYTE_AVAILABLE = True
    OPENAI_AVAILABLE = True
except ImportError:
    # Use environment variables for deployment
    DOMAIN = os.getenv('EGNYTE_DOMAIN')
    CLIENT_ID = os.getenv('EGNYTE_CLIENT_ID')
    CLIENT_SECRET = os.getenv('EGNYTE_CLIENT_SECRET')
    USERNAME = os.getenv('EGNYTE_USERNAME')
    PASSWORD = os.getenv('EGNYTE_PASSWORD')
    ROOT_FOLDER = os.getenv('EGNYTE_ROOT_FOLDER')
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
    EGNYTE_AVAILABLE = all([DOMAIN, CLIENT_ID, CLIENT_SECRET, USERNAME, PASSWORD, ROOT_FOLDER])
    OPENAI_AVAILABLE = bool(OPENAI_API_KEY)

# Rate limiting configuration for Egnyte
# Egnyte limits: 2 calls per second, 1,000 calls per day
RATE_LIMIT_DELAY = 1.0  # Wait 1.0 seconds between calls (allows 1 call/sec, safely under 2/sec)

# Token caching to reduce authentication requests
_egnyte_token_cache = {
    'token': None,
    'expires_at': None
}

# Persistent token storage file
TOKEN_CACHE_FILE = 'egnyte_token_cache.json'

def rate_limit_delay():
    """Add delay to respect rate limits"""
    time.sleep(RATE_LIMIT_DELAY)

def save_token_to_file(token_data):
    """Save token to persistent file"""
    try:
        with open(TOKEN_CACHE_FILE, 'w') as f:
            json.dump(token_data, f)
        logger.info(f"💾 Token saved to {TOKEN_CACHE_FILE}")
    except Exception as e:
        logger.warning(f"Could not save token to file: {e}")

def load_token_from_file():
    """Load token from persistent file"""
    try:
        if os.path.exists(TOKEN_CACHE_FILE):
            with open(TOKEN_CACHE_FILE, 'r') as f:
                token_data = json.load(f)
            logger.info(f"📂 Token loaded from {TOKEN_CACHE_FILE}")
            return token_data
    except Exception as e:
        logger.warning(f"Could not load token from file: {e}")
    return None

def clear_egnyte_token_cache():
    """Clear the cached Egnyte token"""
    global _egnyte_token_cache
    _egnyte_token_cache = {
        'token': None,
        'expires_at': None
    }
    # Also delete the persistent file
    try:
        if os.path.exists(TOKEN_CACHE_FILE):
            os.remove(TOKEN_CACHE_FILE)
            logger.info(f"🗑️ Deleted {TOKEN_CACHE_FILE}")
    except Exception as e:
        logger.warning(f"Could not delete token file: {e}")
    logger.info("🗑️ Cleared Egnyte token cache")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
CORS(app) 

# In-memory storage for job results
job_results = {}
job_status = {}

# Egnyte API Functions
def get_egnyte_token():
    """Get Egnyte access token with rate limiting and retry logic"""
    if not EGNYTE_AVAILABLE:
        logger.error("Egnyte credentials not available")
        return None
    
    # Check if we have a cached token that's still valid
    current_time = time.time()
    if (_egnyte_token_cache['token'] and 
        _egnyte_token_cache['expires_at'] and 
        current_time < _egnyte_token_cache['expires_at']):
        logger.info("✅ Using cached Egnyte token")
        return _egnyte_token_cache['token']
    
    # Try to load from persistent file
    if not _egnyte_token_cache['token']:
        persistent_token = load_token_from_file()
        if persistent_token:
            _egnyte_token_cache.update(persistent_token)
            # Check if the loaded token is still valid
            if (_egnyte_token_cache['token'] and 
                _egnyte_token_cache['expires_at'] and 
                current_time < _egnyte_token_cache['expires_at']):
                logger.info("✅ Using persistent Egnyte token")
                return _egnyte_token_cache['token']
            else:
                logger.info("📂 Persistent token expired, will request new one")
    
    logger.info(f"🔐 Attempting Egnyte authentication...")
    logger.info(f"   Domain: {DOMAIN}")
    logger.info(f"   Username: {USERNAME}")
    logger.info(f"   Client ID: {CLIENT_ID}")
    logger.info(f"   Client Secret: {'*' * len(CLIENT_SECRET) if CLIENT_SECRET else 'None'}")
    logger.info(f"   Password: {'*' * len(PASSWORD) if PASSWORD else 'None'}")
        
    url = f"https://{DOMAIN}/puboauth/token"
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    data = {
        "grant_type": "password",
        "username": USERNAME,
        "password": PASSWORD,
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET
    }
    
    logger.info(f"🌐 Making request to: {url}")
    logger.info(f"📋 Request data: grant_type=password, username={USERNAME}, client_id={CLIENT_ID}, client_secret={'*' * len(CLIENT_SECRET) if CLIENT_SECRET else 'None'}")
    
    max_retries = 3
    retry_count = 0
    
    while retry_count < max_retries:
        try:
            # Add rate limiting delay before each attempt
            rate_limit_delay()
            
            # Use urllib.parse.urlencode to properly encode form data
            encoded_data = urllib.parse.urlencode(data)
            logger.info(f"🔧 Encoded data: {encoded_data.replace(PASSWORD, '*' * len(PASSWORD)).replace(CLIENT_SECRET, '*' * len(CLIENT_SECRET))}")
            
            response = requests.post(url, data=encoded_data, headers=headers)
            
            logger.info(f"📊 Response Status: {response.status_code}")
            logger.info(f"📋 Response Headers: {dict(response.headers)}")
            logger.info(f"📄 Response Body: {response.text}")
            
            if response.status_code == 200:
                token_data = response.json()
                access_token = token_data["access_token"]
                expires_in = token_data.get("expires_in", 3600)  # Default to 1 hour
                token_type = token_data.get("token_type", "unknown")
                
                # Handle token expiration
                if expires_in == -1:
                    # Token never expires, set expiration to a very far future date (10 years)
                    _egnyte_token_cache['expires_at'] = current_time + (10 * 365 * 24 * 3600)  # 10 years from now
                    logger.info(f"   Token never expires (expires_in: -1), setting cache to 10 years")
                else:
                    # Token expires, set expiration 5 minutes early
                    _egnyte_token_cache['expires_at'] = current_time + expires_in - 300
                
                # Cache the token with expiration
                _egnyte_token_cache['token'] = access_token
                
                # Save token persistently
                save_token_to_file(_egnyte_token_cache)
                
                logger.info(f"✅ Authentication successful!")
                logger.info(f"   Token Type: {token_type}")
                logger.info(f"   Expires In: {expires_in} seconds")
                logger.info(f"   Access Token: {access_token[:20]}...")
                logger.info(f"   Token cached until: {datetime.fromtimestamp(_egnyte_token_cache['expires_at'])}")
                
                return access_token
            elif response.status_code == 429:
                # Handle rate limiting
                retry_after = response.headers.get('Retry-After', '30')
                try:
                    retry_seconds = int(retry_after.split(',')[0])  # Handle "2665, 30" format
                except (ValueError, IndexError):
                    retry_seconds = 30
                
                logger.warning(f"XX Rate limit hit! Retry-After: {retry_seconds} seconds")
                logger.warning(f"   Retry-After header: {retry_after}")
                
                # If retry time is too long (>60 seconds), fail immediately to avoid worker timeout
                if retry_seconds > 60:
                    logger.error(f"XX Rate limit retry time ({retry_seconds}s) exceeds worker timeout. Failing immediately.")
                    logger.error("💡 Consider implementing background job processing for bulk requests.")
                    return None
                
                # Only wait for short delays
                logger.warning(f"⚠️ Waiting {retry_seconds} seconds before retry {retry_count + 1}/{max_retries}")
                time.sleep(retry_seconds)
                retry_count += 1
                continue
            else:
                logger.error(f"XX Authentication failed!")
                logger.error(f"   Status Code: {response.status_code}")
                logger.error(f"   Response Text: {response.text}")
                
                # Check for specific error types
                error_text = response.text.lower()
                if "invalid username" in error_text or "invalid password" in error_text:
                    logger.error(f"   Error Type: Invalid credentials")
                elif "rate limit" in error_text or "qps" in error_text:
                    logger.error(f"   Error Type: Rate limiting")
                elif "ip" in error_text or "restricted" in error_text:
                    logger.error(f"   Error Type: IP restriction")
                elif "locked" in error_text or "suspended" in error_text:
                    logger.error(f"   Error Type: Account locked/suspended")
                else:
                    logger.error(f"   Error Type: Unknown")
                
                # Check for specific headers that might indicate the issue
                if 'X-Mashery-Error-Code' in response.headers:
                    error_code = response.headers['X-Mashery-Error-Code']
                    logger.error(f"   Mashery Error Code: {error_code}")
                
                if 'Retry-After' in response.headers:
                    retry_after = response.headers['Retry-After']
                    logger.error(f"   Retry After: {retry_after} seconds")
                
                # For non-429 errors, don't retry
                return None
                
        except requests.exceptions.ConnectionError as e:
            logger.error(f"XX Connection error: {e}")
            logger.error(f"   This might indicate network issues or domain problems")
            return None
        except requests.exceptions.Timeout as e:
            logger.error(f"XX Timeout error: {e}")
            logger.error(f"   Request timed out - server might be slow or unreachable")
            return None
        except requests.exceptions.RequestException as e:
            logger.error(f"XX Request error: {e}")
            logger.error(f"   General request failure")
            return None
        except Exception as e:
            logger.error(f"XX Unexpected error getting Egnyte token: {e}")
            logger.error(f"   Error type: {type(e).__name__}")
            return None
    
    # If we've exhausted all retries
    logger.error(f"XX Authentication failed after {max_retries} retries due to rate limiting")
    return None

def create_egnyte_folder(access_token, parent_folder_id, folder_name):
    """Create a new folder in Egnyte"""
    url = f"https://{DOMAIN}/pubapi/v1/fs/ids/folder/{parent_folder_id}"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    data = {
        "action": "add_folder",
        "name": folder_name
    }
    
    try:
        # Add rate limiting delay
        rate_limit_delay()
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()
    except requests.HTTPError as e:
        error_text = e.response.text.lower()
        
        # Check for folder already exists error
        if "already exists" in error_text or "duplicate" in error_text or "exists" in error_text:
            logger.info(f"Folder '{folder_name}' already exists, checking for existing folder...")
            # Try to find the existing folder
            existing_folder = find_existing_folder(access_token, parent_folder_id, folder_name)
            if existing_folder:
                logger.info(f"Found existing folder: {existing_folder.get('name')}")
                return existing_folder
            else:
                logger.error(f"Folder '{folder_name}' already exists but couldn't find it")
                return None
        
        # Check for rate limiting
        elif "developer over qps" in error_text or "rate limit" in error_text or "qps" in error_text:
            logger.warning(f"Rate limit hit, waiting 3 seconds before retry...")
            time.sleep(3)  # Longer wait for rate limit recovery
            try:
                rate_limit_delay()  # Add rate limiting before retry
                response = requests.post(url, headers=headers, json=data)
                response.raise_for_status()
                return response.json()
            except requests.HTTPError as retry_e:
                logger.error(f"Failed to create folder on retry: {retry_e}")
                return None
        
        # Check for permission errors
        elif e.response.status_code == 403:
            logger.error(f"Permission denied creating folder '{folder_name}': {e.response.text}")
            return None
        
        else:
            logger.error(f"Failed to create folder '{folder_name}': {e.response.text}")
            return None
    except Exception as e:
        logger.error(f"Error creating Egnyte folder '{folder_name}': {e}")
        return None

def find_existing_folder(access_token, parent_folder_id, folder_name):
    """Find an existing folder by name in the parent folder"""
    try:
        folder_data = list_egnyte_folder_contents(access_token, parent_folder_id)
        if folder_data and isinstance(folder_data, dict):
            folders = folder_data.get("folders", [])
            for folder in folders:
                if folder.get('name') == folder_name:
                    return folder
        return None
    except Exception as e:
        logger.error(f"Error finding existing folder: {e}")
        return None

def get_egnyte_folder_details(access_token, folder_id):
    """Get folder details"""
    url = f"https://{DOMAIN}/pubapi/v1/fs/ids/folder/{folder_id}"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    try:
        # Add rate limiting delay
        rate_limit_delay()
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()
    except requests.HTTPError as e:
        if "Developer Over Qps" in e.response.text:
            logger.warning(f"Rate limit hit, waiting 3 seconds before retry...")
            time.sleep(3)  # Longer wait for rate limit recovery
            try:
                rate_limit_delay()  # Add rate limiting before retry
                response = requests.get(url, headers=headers)
                response.raise_for_status()
                return response.json()
            except requests.HTTPError as retry_e:
                logger.error(f"Failed to get folder details on retry: {retry_e}")
                return None
        else:
            logger.error(f"Failed to get folder details: {e}")
            return None
    except Exception as e:
        logger.error(f"Error getting Egnyte folder details: {e}")
        return None

def list_egnyte_folder_contents(access_token, folder_id):
    """List folder contents"""
    url = f"https://{DOMAIN}/pubapi/v1/fs/ids/folder/{folder_id}"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    params = {
        "list_content": "true",
        "count": "100"
    }
    
    try:
        # Add rate limiting delay
        rate_limit_delay()
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        return response.json()
    except requests.HTTPError as e:
        if "Developer Over Qps" in e.response.text:
            logger.warning(f"Rate limit hit, waiting 3 seconds before retry...")
            time.sleep(3)  # Longer wait for rate limit recovery
            try:
                rate_limit_delay()  # Add rate limiting before retry
                response = requests.get(url, headers=headers, params=params)
                response.raise_for_status()
                return response.json()
            except requests.HTTPError as retry_e:
                logger.error(f"Failed to list folder contents on retry: {retry_e}")
                return None
        else:
            logger.error(f"Failed to list folder contents: {e}")
            return None
    except Exception as e:
        logger.error(f"Error listing Egnyte folder contents: {e}")
        return None

def background_create_egnyte_folders(molecule_code: str, campaign_number: str):
    """Background function to create Egnyte folder structure"""
    job_key = f"egnyte_{molecule_code}_{campaign_number}"
    
    logger.info(f"BACKGROUND: Starting Egnyte folder creation for {job_key}")
    
    try:
        # Update status to running
        job_status[job_key] = {
            "status": "running",
            "message": "Creating Egnyte folder structure...",
            "started_at": datetime.now().isoformat(),
            "progress": 0
        }
        
        # Get access token
        access_token = get_egnyte_token()
        if not access_token:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to get Egnyte access token",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Update progress
        job_status[job_key]["progress"] = 20
        job_status[job_key]["message"] = "Creating project folder..."
        
        # Check if project folder already exists
        project_folder_name = f"Project; Molecule {molecule_code}"
        
        # First, try to list existing folders to check for duplicates
        existing_folders_data = list_egnyte_folder_contents(access_token, ROOT_FOLDER)
        if existing_folders_data and isinstance(existing_folders_data, dict):
            existing_folders = existing_folders_data.get("folders", [])
            for folder in existing_folders:
                if folder.get('name') == project_folder_name:
                    # Check if campaign folder already exists
                    campaign_folder_name = f"Project {molecule_code} (Campaign #{campaign_number})"
                    campaign_contents_data = list_egnyte_folder_contents(access_token, folder.get('folder_id'))
                    if campaign_contents_data and isinstance(campaign_contents_data, dict):
                        campaign_folders = campaign_contents_data.get("folders", [])
                        for campaign_folder in campaign_folders:
                            if campaign_folder.get('name') == campaign_folder_name:
                                job_status[job_key] = {
                                    "status": "failed",
                                    "message": f"Project {molecule_code} Campaign {campaign_number} already exists",
                                    "started_at": job_status[job_key]["started_at"],
                                    "completed_at": datetime.now().isoformat()
                                }
                                return
        
        # Create project folder
        project_folder = create_egnyte_folder(access_token, ROOT_FOLDER, project_folder_name)
        
        if not project_folder or not isinstance(project_folder, dict):
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to create project folder",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        project_folder_id = project_folder.get('folder_id')
        
        # Update progress
        job_status[job_key]["progress"] = 40
        job_status[job_key]["message"] = "Creating campaign folder..."
        
        # Create campaign folder
        campaign_folder_name = f"Project {molecule_code} (Campaign #{campaign_number})"
        campaign_folder = create_egnyte_folder(access_token, project_folder_id, campaign_folder_name)
        
        if not campaign_folder or not isinstance(campaign_folder, dict):
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to create campaign folder",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        campaign_folder_id = campaign_folder.get('folder_id')
        
        # Update progress
        job_status[job_key]["progress"] = 60
        job_status[job_key]["message"] = "Creating Pre and Post folders..."
        
        # Create Pre and Post folders
        pre_folder = create_egnyte_folder(access_token, campaign_folder_id, "Pre")
        post_folder = create_egnyte_folder(access_token, campaign_folder_id, "Post")
        
        if not pre_folder or not isinstance(pre_folder, dict) or not post_folder or not isinstance(post_folder, dict):
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to create Pre/Post folders",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Create department folders under Pre and Post
        departments = ["mfg", "Anal", "Stability", "CTM"]
        statuses = ["Draft", "Review", "Approved"]
        
        job_status[job_key]["progress"] = 65
        job_status[job_key]["message"] = "Creating department folders..."
        
        for phase_name, phase_folder in [("Pre", pre_folder), ("Post", post_folder)]:
            phase_folder_id = phase_folder.get('folder_id')
            
            for dept in departments:
                dept_folder = create_egnyte_folder(access_token, phase_folder_id, dept)
                if dept_folder and isinstance(dept_folder, dict):
                    dept_folder_id = dept_folder.get('folder_id')
                    
                    # Create status folders under each department
                    for status in statuses:
                        create_egnyte_folder(access_token, dept_folder_id, status)
        
        # Update progress
        job_status[job_key]["progress"] = 80
        job_status[job_key]["message"] = "Creating Draft AI Reg Document folder..."
        
        # Create Draft AI Reg Document folder under project
        reg_doc_folder = create_egnyte_folder(access_token, project_folder_id, "Draft AI Reg Document")
        if not reg_doc_folder or not isinstance(reg_doc_folder, dict):
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to create Draft AI Reg Document folder",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        reg_doc_folder_id = reg_doc_folder.get('folder_id')
        
        # Update progress
        job_status[job_key]["progress"] = 85
        job_status[job_key]["message"] = "Creating regulatory document folders..."
        
        # Create regulatory document folders under Draft AI Reg Document
        reg_types = ["IND", "IMPD", "Canada"]
        for reg_type in reg_types:
            reg_type_folder = create_egnyte_folder(access_token, reg_doc_folder_id, reg_type)
            if reg_type_folder and isinstance(reg_type_folder, dict):
                reg_type_folder_id = reg_type_folder.get('folder_id')
                
                for status in statuses:
                    create_egnyte_folder(access_token, reg_type_folder_id, status)
        
        # Store the result with URLs
        job_results[job_key] = {
            "project_folder": {
                "id": project_folder.get('folder_id'),
                "name": project_folder.get('name'),
                "path": project_folder.get('path'),
                "url": f"https://{DOMAIN}/app/index.do#storage/files/1{project_folder.get('path', '')}"
            },
            "campaign_folder": {
                "id": campaign_folder.get('folder_id'),
                "name": campaign_folder.get('name'),
                "path": campaign_folder.get('path'),
                "url": f"https://{DOMAIN}/app/index.do#storage/files/1{campaign_folder.get('path', '')}"
            },
            "reg_doc_folder": {
                "id": reg_doc_folder.get('folder_id'),
                "name": reg_doc_folder.get('name'),
                "path": reg_doc_folder.get('path'),
                "url": f"https://{DOMAIN}/app/index.do#storage/files/1{reg_doc_folder.get('path', '')}"
            },
            "molecule_code": molecule_code,
            "campaign_number": campaign_number,
            "folder_structure": {
                "project_name": project_folder.get('name'),
                "campaign_name": campaign_folder.get('name'),
                "reg_doc_name": reg_doc_folder.get('name'),
                "project_url": f"https://{DOMAIN}/app/index.do#storage/files/1{project_folder.get('path', '')}",
                "campaign_url": f"https://{DOMAIN}/app/index.do#storage/files/1{campaign_folder.get('path', '')}",
                "reg_doc_url": f"https://{DOMAIN}/app/index.do#storage/files/1{reg_doc_folder.get('path', '')}"
            }
        }
        
        # Update status to completed
        job_status[job_key] = {
            "status": "completed",
            "message": "Egnyte folder structure created successfully",
            "started_at": job_status[job_key]["started_at"],
            "completed_at": datetime.now().isoformat(),
            "progress": 100
        }
        
        logger.info(f"Background Egnyte job completed for {job_key}")
        
    except Exception as e:
        logger.error(f"Background Egnyte job failed for {job_key}: {e}")
        job_status[job_key] = {
            "status": "failed",
            "message": f"Error: {str(e)}",
            "started_at": job_status[job_key]["started_at"],
            "completed_at": datetime.now().isoformat()
        }

def load_openai_api_key():
    """Load OpenAI API key from credentials or environment variable"""
    return OPENAI_API_KEY

def initialize_openai():
    """Initialize OpenAI client"""
    if not OPENAI_AVAILABLE:
        logger.error("OpenAI credentials not available")
        return None
    
    api_key = load_openai_api_key()
    if api_key and api_key != "your-openai-api-key-here":
        try:
            client = OpenAI(api_key=api_key)
            return client
        except Exception as e:
            logger.error(f"Error initializing OpenAI client: {e}")
            return None
    else:
        logger.error("OpenAI API key not found or invalid")
        return None

def create_sample_pharma_data():
    """Create sample pharmaceutical composition data"""
    data = {
        'Component': [
            'Active Pharmaceutical Ingredient',
            'Microcrystalline Cellulose',
            'Lactose Monohydrate',
            'Croscarmellose Sodium',
            'Magnesium Stearate',
            'Opadry II White'
        ],
        'Quality_Reference': ['USP', 'NF', 'NF', 'NF', 'NF', 'NF'],
        'Function': [
            'Active Ingredient',
            'Tablet Diluent',
            'Tablet Diluent',
            'Disintegrant',
            'Lubricant',
            'Film Coating'
        ],
        'Quantity_mg_per_tablet': [25.0, 150.0, 100.0, 10.0, 2.0, 8.0]
    }
    return pd.DataFrame(data)

# Legacy functions removed - replaced with upload_files_prompt_to_openai

def export_to_pdf_regulatory(df: pd.DataFrame, sections: Dict[str, str], 
                            product_code: str, dosage_form: str,
                            molecule_code: str = None, campaign_number: str = None):
    """Export regulatory document to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Add title with campaign information
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1
    )
    title_text = 'Section 3.2.P.1 Description and Composition of the Drug Product'
    if molecule_code and campaign_number:
        title_text += f' - Molecule {molecule_code} Campaign {campaign_number}'
    story.append(Paragraph(title_text, title_style))
    story.append(Spacer(1, 12))
    
    # Add campaign information header
    if molecule_code and campaign_number:
        story.append(Paragraph('Project Information', styles['Heading2']))
        story.append(Paragraph(f'Molecule Code: {molecule_code}', styles['Normal']))
        story.append(Paragraph(f'Campaign Number: {campaign_number}', styles['Normal']))
        story.append(Paragraph(f'Product Code: {product_code}', styles['Normal']))
        story.append(Paragraph(f'Dosage Form: {dosage_form}', styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add description section
    story.append(Paragraph('3.2.P.1.1 Description of the Dosage Form', styles['Heading2']))
    story.append(Paragraph(sections['description'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Add composition section
    story.append(Paragraph('3.2.P.1.2 Composition', styles['Heading2']))
    story.append(Paragraph(sections['composition_intro'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Prepare table data
    headers = ['Component', 'Quality Reference', 'Function', 'Quantity / Unit (mg per tablet)']
    table_data = [headers]
    
    total_weight = 0
    for _, row in df.iterrows():
        table_data.append([
            str(row['Component']),
            str(row['Quality_Reference']),
            str(row['Function']),
            str(row['Quantity_mg_per_tablet'])
        ])
        total_weight += row['Quantity_mg_per_tablet']
    
    # Add total weight row
    table_data.append(['Total Weight', '', '', str(total_weight)])
    
    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    story.append(table)
    story.append(Spacer(1, 12))
    
    # Add footnote
    story.append(Paragraph("Abbreviations: NF = National Formulary; Ph. Eur. = European Pharmacopoeia; USP = United States Pharmacopoeia.", styles['Normal']))
    
    # Add pharmaceutical development
    if sections.get('pharmaceutical_development'):
        story.append(Paragraph('3.2.P.1.3 Pharmaceutical Development', styles['Heading2']))
        story.append(Paragraph(sections['pharmaceutical_development'], styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add manufacturing process
    if sections.get('manufacturing_process'):
        story.append(Paragraph('3.2.P.1.4 Manufacturing Process', styles['Heading2']))
        story.append(Paragraph(sections['manufacturing_process'], styles['Normal']))
        story.append(Spacer(1, 12))
    
    doc.build(story)
    return buffer


@app.route('/test-thread', methods=['GET'])
def test_thread():
    """Test endpoint to verify threading works"""
    import time
    
    def background_test():
        logger.info("BACKGROUND TEST: Starting sleep...")
        time.sleep(10)  # Sleep for 10 seconds
        logger.info("BACKGROUND TEST: Sleep completed!")
    
    logger.info("TEST: Starting background thread test")
    thread = threading.Thread(target=background_test, daemon=True)
    thread.start()
    logger.info("TEST: Background thread started, returning immediately")
    
    return jsonify({
        "status": "success",
        "message": "Background thread test started",
        "timestamp": datetime.now().isoformat()
    })

@app.route('/folder-status', methods=['GET'])
def folder_status():
    """Check the status of a folder creation job"""
    try:
        molecule_code = request.args.get('molecule_code')
        campaign_number = request.args.get('campaign_number')
        
        if not molecule_code or not campaign_number:
            return jsonify({"error": "molecule_code and campaign_number are required"}), 400
        
        job_key = f"{molecule_code}_{campaign_number}"
        
        if job_key not in job_status:
            return jsonify({
                "status": "not_found",
                "message": "No job found for this molecule and campaign",
                "job_key": job_key
            })
        
        status_info = job_status[job_key]
        
        response = {
            "status": status_info["status"],
            "message": status_info["message"],
            "job_key": job_key,
            "started_at": status_info["started_at"]
        }
        
        # Add progress if available
        if "progress" in status_info:
            response["progress"] = status_info["progress"]
        
        # Add completion time if available
        if "completed_at" in status_info:
            response["completed_at"] = status_info["completed_at"]
        
        # Add result data if completed
        if status_info["status"] == "completed" and job_key in job_results:
            response["data"] = job_results[job_key]
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in folder_status: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/document-status', methods=['GET'])
def document_status():
    """Check the status of a document generation job"""
    try:
        molecule_code = request.args.get('molecule_code')
        campaign_number = request.args.get('campaign_number')
        
        if not molecule_code or not campaign_number:
            return jsonify({"error": "molecule_code and campaign_number are required"}), 400
        
        job_key = f"doc_{molecule_code}_{campaign_number}"
        
        if job_key not in job_status:
            return jsonify({
                "status": "not_found",
                "message": "No document generation job found for this molecule and campaign",
                "job_key": job_key
            })
        
        status_info = job_status[job_key]
        
        response = {
            "status": status_info["status"],
            "message": status_info["message"],
            "job_key": job_key,
            "started_at": status_info["started_at"]
        }
        
        # Add progress if available
        if "progress" in status_info:
            response["progress"] = status_info["progress"]
        
        # Add completion time if available
        if "completed_at" in status_info:
            response["completed_at"] = status_info["completed_at"]
        
        # Add result data if completed
        if status_info["status"] == "completed" and job_key in job_results:
            response["data"] = job_results[job_key]
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in document_status: {e}")
        return jsonify({"error": str(e)}), 500

# Egnyte API Routes
@app.route('/egnyte-generate-folder-structure', methods=['POST'])
def egnyte_generate_folder_structure():
    """Start background job to generate Egnyte folder structure"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        molecule_code = data.get('molecule_code')
        campaign_number = data.get('campaign_number')
        
        if not molecule_code or not campaign_number:
            return jsonify({"error": "molecule_code and campaign_number are required"}), 400
        
        job_key = f"egnyte_{molecule_code}_{campaign_number}"
        
        # Check if job is already running
        if job_key in job_status and job_status[job_key]["status"] == "running":
            return jsonify({
                "status": "already_running",
                "message": "Egnyte folder creation job is already running",
                "job_key": job_key
            })
        
        # Check if job is already completed
        if job_key in job_status and job_status[job_key]["status"] == "completed":
            return jsonify({
                "status": "already_completed",
                "message": "Egnyte folder structure already exists",
                "job_key": job_key,
                "data": job_results.get(job_key)
            })
        
        # Start background thread
        thread = threading.Thread(
            target=background_create_egnyte_folders,
            args=(molecule_code, campaign_number),
            daemon=True
        )
        thread.start()
        
        return jsonify({
            "status": "started",
            "message": "Egnyte folder creation job started",
            "job_key": job_key,
            "poll_url": f"/egnyte-folder-status?molecule_code={molecule_code}&campaign_number={campaign_number}"
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-folder-status', methods=['GET'])
def egnyte_folder_status():
    """Check the status of an Egnyte folder creation job"""
    try:
        molecule_code = request.args.get('molecule_code')
        campaign_number = request.args.get('campaign_number')
        
        if not molecule_code or not campaign_number:
            return jsonify({"error": "molecule_code and campaign_number are required"}), 400
        
        job_key = f"egnyte_{molecule_code}_{campaign_number}"
        
        if job_key not in job_status:
            return jsonify({
                "status": "not_found",
                "message": "No Egnyte job found for this molecule and campaign",
                "job_key": job_key
            })
        
        status_info = job_status[job_key]
        
        response = {
            "status": status_info["status"],
            "message": status_info["message"],
            "job_key": job_key,
            "started_at": status_info["started_at"]
        }
        
        # Add progress if available
        if "progress" in status_info:
            response["progress"] = status_info["progress"]
        
        # Add completion time if available
        if "completed_at" in status_info:
            response["completed_at"] = status_info["completed_at"]
        
        # Add result data if completed
        if status_info["status"] == "completed" and job_key in job_results:
            response["data"] = job_results[job_key]
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-list-folder', methods=['GET'])
def egnyte_list_folder():
    """List contents of an Egnyte folder"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        folder_id = request.args.get('folder_id', ROOT_FOLDER)
        
        access_token = get_egnyte_token()
        if not access_token:
            return jsonify({"error": "Failed to get Egnyte access token"}), 500
        
        folder_data = list_egnyte_folder_contents(access_token, folder_id)
        if not folder_data:
            return jsonify({"error": "Failed to list folder contents"}), 500
        
        return jsonify({
            "status": "success",
            "folder_id": folder_id,
            "folders": folder_data.get("folders", []),
            "files": folder_data.get("files", []),
            "total_folders": len(folder_data.get("folders", [])),
            "total_files": len(folder_data.get("files", []))
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-create-folder', methods=['POST'])
def egnyte_create_folder():
    """Create a new folder in Egnyte"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        parent_folder_id = data.get('parent_folder_id', ROOT_FOLDER)
        folder_name = data.get('folder_name')
        
        if not folder_name:
            return jsonify({"error": "folder_name is required"}), 400
        
        access_token = get_egnyte_token()
        if not access_token:
            return jsonify({"error": "Failed to get Egnyte access token"}), 500
        
        result = create_egnyte_folder(access_token, parent_folder_id, folder_name)
        if not result:
            return jsonify({"error": "Failed to create folder"}), 500
        
        return jsonify({
            "status": "success",
            "message": "Folder created successfully",
            "data": {
                "folder_id": result.get('folder_id'),
                "name": result.get('name'),
                "path": result.get('path'),
                "url": f"https://{DOMAIN}/app/index.do#storage/files/1{result.get('path', '')}"
            }
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-list-templates', methods=['GET'])
def egnyte_list_templates():
    """List all items in the templates folder"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        templates_folder_id = "966281ab-54c3-47ea-b20f-b38ed2ef9b30"
        
        access_token = get_egnyte_token()
        if not access_token:
            return jsonify({"error": "Failed to get Egnyte access token"}), 500
        
        folder_data = list_egnyte_folder_contents(access_token, templates_folder_id)
        if not folder_data:
            return jsonify({"error": "Failed to list templates folder contents"}), 500
        
        return jsonify({
            "status": "success",
            "folder_id": templates_folder_id,
            "templates": folder_data.get("files", []),
            "folders": folder_data.get("folders", []),
            "total_templates": len(folder_data.get("files", [])),
            "total_folders": len(folder_data.get("folders", []))
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-list-source-documents', methods=['GET'])
def egnyte_list_source_documents():
    """List all items in the source documents folder"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        source_docs_folder_id = "56545792-6b5d-4fc3-8e78-31d401bd7088"
        
        access_token = get_egnyte_token()
        if not access_token:
            return jsonify({"error": "Failed to get Egnyte access token"}), 500
        
        folder_data = list_egnyte_folder_contents(access_token, source_docs_folder_id)
        if not folder_data:
            return jsonify({"error": "Failed to list source documents folder contents"}), 500
        
        return jsonify({
            "status": "success",
            "folder_id": source_docs_folder_id,
            "documents": folder_data.get("files", []),
            "folders": folder_data.get("folders", []),
            "total_documents": len(folder_data.get("files", [])),
            "total_folders": len(folder_data.get("folders", []))
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-download-file', methods=['POST'])
def egnyte_download_file():
    """Download a file from Egnyte"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({"error": "file_id is required"}), 400
        
        access_token = get_egnyte_token()
        if not access_token:
            return jsonify({"error": "Failed to get Egnyte access token"}), 500
        
        # Download the file
        url = f"https://{DOMAIN}/pubapi/v1/fs-content/ids/file/{file_id}"
        headers = {
            "Authorization": f"Bearer {access_token}"
        }
        
        try:
            rate_limit_delay()
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            # Return the file content as base64
            file_content = base64.b64encode(response.content).decode('utf-8')
            
            return jsonify({
                "status": "success",
                "file_id": file_id,
                "content": file_content,
                "content_type": response.headers.get('content-type', 'application/octet-stream'),
                "size": len(response.content)
            })
            
        except requests.HTTPError as e:
            logger.error(f"Failed to download file {file_id}: {e}")
            return jsonify({"error": f"Failed to download file: {e}"}), 500
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-generate-document', methods=['POST'])
def egnyte_generate_document():
    """Generate a new document using OpenAI and save to Egnyte"""
    if not EGNYTE_AVAILABLE:
        return jsonify({"error": "Egnyte integration not available. Please configure Egnyte credentials."}), 503
        
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        template_file_id = data.get('template_file_id')
        source_document_ids = data.get('source_document_ids', [])
        molecule_code = data.get('molecule_code')
        campaign_number = data.get('campaign_number')
        document_name = data.get('document_name', 'Generated Document')
        
        if not template_file_id:
            return jsonify({"error": "template_file_id is required"}), 400
        
        if not source_document_ids:
            return jsonify({"error": "source_document_ids is required"}), 400
        
        if not molecule_code or not campaign_number:
            return jsonify({"error": "molecule_code and campaign_number are required"}), 400
        
        # Create job key for tracking
        job_key = f"doc_gen_{molecule_code}_{campaign_number}_{int(time.time())}"
        
        # Start background thread
        thread = threading.Thread(
            target=background_generate_egnyte_document,
            args=(template_file_id, source_document_ids, molecule_code, campaign_number, document_name, job_key),
            daemon=True
        )
        thread.start()
        
        return jsonify({
            "status": "started",
            "message": "Document generation job started",
            "job_key": job_key,
            "poll_url": f"/egnyte-document-status?job_key={job_key}"
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/egnyte-document-status', methods=['GET'])
def egnyte_document_status():
    """Check the status of a document generation job"""
    try:
        job_key = request.args.get('job_key')
        
        if not job_key:
            return jsonify({"error": "job_key is required"}), 400
        
        if job_key not in job_status:
            return jsonify({
                "status": "not_found",
                "message": "No document generation job found",
                "job_key": job_key
            })
        
        status_info = job_status[job_key]
        
        response = {
            "status": status_info["status"],
            "message": status_info["message"],
            "job_key": job_key,
            "started_at": status_info["started_at"]
        }
        
        # Add progress if available
        if "progress" in status_info:
            response["progress"] = status_info["progress"]
        
        # Add completion time if available
        if "completed_at" in status_info:
            response["completed_at"] = status_info["completed_at"]
        
        # Add result data if completed
        if status_info["status"] == "completed" and job_key in job_results:
            response["data"] = job_results[job_key]
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def background_generate_egnyte_document(template_file_id, source_document_ids, molecule_code, campaign_number, document_name, job_key):
    """Background function to generate document using OpenAI and save to Egnyte"""
    logger.info(f"BACKGROUND: Starting document generation for {job_key}")
    
    try:
        # Update status to running
        job_status[job_key] = {
            "status": "running",
            "message": "Starting document generation...",
            "started_at": datetime.now().isoformat(),
            "progress": 0
        }
        
        # Get access token
        access_token = get_egnyte_token()
        if not access_token:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to get Egnyte access token",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Update progress
        job_status[job_key]["progress"] = 10
        job_status[job_key]["message"] = "Downloading template file..."
        
        # Download template file
        template_content = download_egnyte_file(access_token, template_file_id)
        if not template_content:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to download template file",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Update progress
        job_status[job_key]["progress"] = 20
        job_status[job_key]["message"] = "Downloading source documents..."
        
        # Download source documents
        source_contents = []
        for i, doc_id in enumerate(source_document_ids):
            content = download_egnyte_file(access_token, doc_id)
            if content:
                source_contents.append(content)
            
            # Update progress for each document
            progress = 20 + (i + 1) * 20 // len(source_document_ids)
            job_status[job_key]["progress"] = progress
            job_status[job_key]["message"] = f"Downloaded {i + 1}/{len(source_document_ids)} source documents..."
        
        if not source_contents:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to download any source documents",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Update progress
        job_status[job_key]["progress"] = 60
        job_status[job_key]["message"] = "Generating document with OpenAI..."
        
        # Generate document using OpenAI
        generated_content = generate_document_with_openai(template_content, source_contents, document_name)
        if not generated_content:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to generate document with OpenAI",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Update progress
        job_status[job_key]["progress"] = 80
        job_status[job_key]["message"] = "Saving document to Egnyte..."
        
        # Find the target folder in Egnyte
        target_folder_id = find_egnyte_target_folder(access_token, molecule_code, campaign_number)
        if not target_folder_id:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to find target folder in Egnyte",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Upload the generated document
        file_name = f"{document_name}_{molecule_code}_Campaign_{campaign_number}.docx"
        uploaded_file = upload_file_to_egnyte(access_token, target_folder_id, file_name, generated_content)
        if not uploaded_file:
            job_status[job_key] = {
                "status": "failed",
                "message": "Failed to upload document to Egnyte",
                "started_at": job_status[job_key]["started_at"],
                "completed_at": datetime.now().isoformat()
            }
            return
        
        # Store the result
        job_results[job_key] = {
            "status": "success",
            "file_name": file_name,
            "file_id": uploaded_file.get('entry_id'),
            "file_url": f"https://{DOMAIN}/app/index.do#storage/files/1{uploaded_file.get('path', '')}",
            "molecule_code": molecule_code,
            "campaign_number": campaign_number,
            "document_name": document_name
        }
        
        # Update status to completed
        job_status[job_key] = {
            "status": "completed",
            "message": "Document generated and saved successfully",
            "started_at": job_status[job_key]["started_at"],
            "completed_at": datetime.now().isoformat(),
            "progress": 100
        }
        
        logger.info(f"Background document generation completed for {job_key}")
        
    except Exception as e:
        logger.error(f"Background document generation failed for {job_key}: {e}")
        job_status[job_key] = {
            "status": "failed",
            "message": f"Error: {str(e)}",
            "started_at": job_status[job_key]["started_at"],
            "completed_at": datetime.now().isoformat()
        }

def download_egnyte_file(access_token, file_id, file_path=None):
    """Download a file from Egnyte and return its content"""
    try:
        logger.info(f"Downloading Egnyte file: {file_id}")
        
        # Try path-based download first if path is available
        if file_path:
            logger.info(f"Attempting path-based download: {file_path}")
            url = f"https://{DOMAIN}/pubapi/v1/fs-content{file_path}"
            headers = {
                "Authorization": f"Bearer {access_token}"
            }
            
            logger.info(f"Egnyte URL (path-based): {url}")
            logger.info(f"Domain: {DOMAIN}")
            logger.info(f"Access token length: {len(access_token) if access_token else 0}")
            
            rate_limit_delay()
            logger.info("Making path-based request to Egnyte API...")
            response = requests.get(url, headers=headers)
            
            logger.info(f"Path-based response status code: {response.status_code}")
            
            if response.status_code == 200:
                content = response.content
                logger.info(f"SUCCESS: Downloaded {len(content)} bytes from Egnyte (path-based)")
                return content
            else:
                logger.warning(f"Path-based download failed with status {response.status_code}, trying ID-based download...")
        
        # Fall back to ID-based download
        logger.info(f"Attempting ID-based download: {file_id}")
        url = f"https://{DOMAIN}/pubapi/v1/fs-content/ids/file/{file_id}"
        headers = {
            "Authorization": f"Bearer {access_token}"
        }
        
        logger.info(f"Egnyte URL (ID-based): {url}")
        logger.info(f"Domain: {DOMAIN}")
        logger.info(f"Access token length: {len(access_token) if access_token else 0}")
        
        rate_limit_delay()
        logger.info("Making ID-based request to Egnyte API...")
        response = requests.get(url, headers=headers)
        
        logger.info(f"ID-based response status code: {response.status_code}")
        logger.info(f"Response headers: {dict(response.headers)}")
        
        if response.status_code != 200:
            logger.error(f"Egnyte API returned error status: {response.status_code}")
            logger.error(f"Response text: {response.text}")
            response.raise_for_status()
        
        content = response.content
        logger.info(f"SUCCESS: Downloaded {len(content)} bytes from Egnyte (ID-based)")
        
        return content
        
    except Exception as e:
        logger.error("=" * 50)
        logger.error("EGNYTE FILE DOWNLOAD FAILED")
        logger.error("=" * 50)
        logger.error(f"File ID: {file_id}")
        logger.error(f"File Path: {file_path}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
        logger.error("=" * 50)
        return None

def find_egnyte_target_folder(access_token, molecule_code, campaign_number):
    """Find the target folder in Egnyte for the generated document"""
    try:
        # First, find the project folder
        project_folder_name = f"Project; Molecule {molecule_code}"
        project_folder = find_folder_by_name(access_token, ROOT_FOLDER, project_folder_name)
        
        if not project_folder:
            logger.error(f"Project folder not found: {project_folder_name}")
            return None
        
        # Find the campaign folder
        campaign_folder_name = f"Project {molecule_code} (Campaign #{campaign_number})"
        campaign_folder = find_folder_by_name(access_token, project_folder.get('folder_id'), campaign_folder_name)
        
        if not campaign_folder:
            logger.error(f"Campaign folder not found: {campaign_folder_name}")
            return None
        
        # Find the Draft AI Reg Document folder
        reg_doc_folder = find_folder_by_name(access_token, project_folder.get('folder_id'), "Draft AI Reg Document")
        
        if not reg_doc_folder:
            logger.error("Draft AI Reg Document folder not found")
            return None
        
        return reg_doc_folder.get('folder_id')
        
    except Exception as e:
        logger.error(f"Error finding target folder: {e}")
        return None

def find_folder_by_name(access_token, parent_folder_id, folder_name):
    """Find a folder by name in a parent folder"""
    try:
        folder_data = list_egnyte_folder_contents(access_token, parent_folder_id)
        if folder_data and isinstance(folder_data, dict):
            folders = folder_data.get("folders", [])
            for folder in folders:
                if folder.get('name') == folder_name:
                    return folder
        return None
    except Exception as e:
        logger.error(f"Error finding folder by name: {e}")
        return None

def upload_file_to_egnyte(access_token, folder_id, file_name, file_content):
    """Upload a file to Egnyte"""
    try:
        url = f"https://{DOMAIN}/pubapi/v1/fs-content/ids/folder/{folder_id}"
        
        # Determine content type based on file extension
        if file_name.lower().endswith('.docx'):
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_name.lower().endswith('.pdf'):
            content_type = "application/pdf"
        else:
            content_type = "application/octet-stream"
        
        headers = {
            "Authorization": f"Bearer {access_token}"
        }
        
        files = {
            'file': (file_name, file_content, content_type)
        }
        
        rate_limit_delay()
        logger.info(f"Uploading file '{file_name}' to folder {folder_id}")
        logger.info(f"File size: {len(file_content)} bytes")
        logger.info(f"Content type: {content_type}")
        
        response = requests.post(url, headers=headers, files=files)
        
        logger.info(f"Upload response status: {response.status_code}")
        if response.status_code != 200:
            logger.error(f"Upload failed with status {response.status_code}")
            logger.error(f"Response text: {response.text}")
        
        response.raise_for_status()
        
        result = response.json()
        logger.info(f"Upload successful for file '{file_name}'")
        return result
        
    except Exception as e:
        logger.error(f"Error uploading file to Egnyte: {e}")
        logger.error(f"File: {file_name}, Folder: {folder_id}")
        return None

def generate_document_with_openai(template_content, source_contents, document_name):
    """Generate a document using OpenAI based on template and source documents"""
    try:
        # Initialize OpenAI
        client = initialize_openai()
        if not client:
            logger.error("Failed to initialize OpenAI")
            return None
        
        # Convert template content to text (assuming it's a Word document)
        template_text = extract_text_from_docx(template_content)
        
        # Convert source documents to text
        source_texts = []
        for content in source_contents:
            text = extract_text_from_docx(content)
            if text:
                source_texts.append(text)
        
        # Create the prompt for OpenAI
        prompt = f"""
        You are a document generation assistant. Please create a new document based on the following:

        TEMPLATE DOCUMENT:
        {template_text}

        SOURCE DOCUMENTS:
        {chr(10).join([f"Document {i+1}: {text}" for i, text in enumerate(source_texts)])}

        INSTRUCTIONS:
        1. Use the template document as the structure and format
        2. Incorporate relevant information from the source documents
        3. Create a comprehensive, well-structured document
        4. Maintain professional tone and formatting
        5. Document name: {document_name}

        Please generate the complete document content.
        """
        
        # Call OpenAI
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a professional document generation assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        generated_text = response.choices[0].message.content
        
        # Convert the generated text back to a Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(document_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        paragraphs = generated_text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
        
    except Exception as e:
        logger.error(f"Error generating document with OpenAI: {e}")
        return None

def extract_text_from_docx(docx_content):
    """Extract text from a Word document"""
    try:
        doc = Document(io.BytesIO(docx_content))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())
        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""

@app.route('/egnyte-clear-cache', methods=['POST'])
def egnyte_clear_cache():
    """Clear the Egnyte token cache"""
    try:
        clear_egnyte_token_cache()
        return jsonify({
            "status": "success",
            "message": "Egnyte token cache cleared"
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/test-document-generation', methods=['POST'])
def test_document_generation():
    """Test the new simplified document generation function"""
    try:
        logger.info("=" * 80)
        logger.info("TESTING SIMPLIFIED DOCUMENT GENERATION")
        logger.info("=" * 80)
        
        # Test files
        template_path = "IND_3.2.P.1_Template.docx"
        source_document_path = "THPG001009 Product Code.pdf"
        
        # Check if test files exist
        if not os.path.exists(template_path):
            return jsonify({
                "status": "error", 
                "message": f"Template file not found at {template_path}"
            }), 404
            
        if not os.path.exists(source_document_path):
            return jsonify({
                "status": "error", 
                "message": f"Source document file not found at {source_document_path}"
            }), 404
        from demo_prompt import test_prompt
        # Test prompt
        
        logger.info("Starting document generation test...")
        start_time = time.time()
        
        # Call the new function
        docx_content = upload_files_prompt_to_openai(
            prompt=test_prompt,
            template_path=template_path,
            source_document_path=source_document_path
        )
        
        end_time = time.time()
        processing_time = end_time - start_time
        
        if docx_content:
            # Save the generated document for inspection
            output_filename = f"generated_document_{int(time.time())}.docx"
            with open(output_filename, 'wb') as f:
                f.write(docx_content)
            
            logger.info(f"✅ Document generation successful!")
            logger.info(f"   Processing time: {processing_time:.2f} seconds")
            logger.info(f"   Output file: {output_filename}")
            logger.info(f"   Document size: {len(docx_content)} bytes")
            
            return jsonify({
                "status": "success",
                "message": "Document generation test completed successfully",
                "processing_time_seconds": processing_time,
                "document_size_bytes": len(docx_content),
                "output_filename": output_filename,
                "details": {
                    "template_file": template_path,
                    "source_document": source_document_path,
                    "prompt_length": len(test_prompt)
                }
            })
        else:
            logger.error("XX Document generation failed - returned None")
            return jsonify({
                "status": "error",
                "message": "Document generation failed - function returned None",
                "processing_time_seconds": processing_time
            }), 500

    except Exception as e:
        logger.error("=" * 80)
        logger.error("DOCUMENT GENERATION TEST FAILED")
        logger.error("=" * 80)
        logger.error(f"Error: {e}")
        import traceback
        logger.error(f"Traceback:\n{traceback.format_exc()}")
        
        return jsonify({
            "status": "error",
            "message": "Document generation test failed",
            "error": str(e)
        }), 500

@app.route('/reg-docs-bulk-request', methods=['POST'])
def reg_docs_bulk_request():
    """Bulk request for regulatory documents"""
    try:
        data = request.get_json()
        timestamp_clean = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

        request_df = pd.DataFrame(data)
        print("Request DataFrame columns:", request_df.columns.tolist())
        
        # Check that request contains docs 
        if len(request_df) == 0:
            return jsonify({"status": "error", "message": "No requests from campaign folder received"}), 400
        
        print("Received bulk request from Retool at ", timestamp_clean)
        print(f"Processing {len(request_df)} requests")
        
        # Get unique product codes
        unique_product_codes = request_df['product_code'].unique().tolist()
        print(f"Unique product codes: {unique_product_codes}")
        
        # Extract root names and latest versions from reg doc versions
        # Handle null values safely
        def extract_root_and_latest(version_str):
            if not version_str or not isinstance(version_str, str):
                return None, None
            
            # Extract root (everything before "v")
            parts = version_str.split('v')
            root = parts[0].strip() if parts[0] else None
            
            # Find all version numbers in the string (e.g., v1.0, v2.0, v3.1)
            import re
            versions = re.findall(r'v(\d+\.?\d*)', version_str)
            
            if not versions:
                return root, None
            
            # Convert to float for comparison and find the latest
            try:
                latest_version = max(float(v) for v in versions)
                latest = f"v{latest_version}"
                return root, latest
            except ValueError:
                return root, None
        
        # Extract both root and latest version
        request_df[['reg_doc_version_active_root', 'reg_doc_version_active_latest']] = pd.DataFrame(
            request_df['reg_doc_version_active'].apply(extract_root_and_latest).tolist(), 
            index=request_df.index
        )
        request_df[['reg_doc_version_placebo_root', 'reg_doc_version_placebo_latest']] = pd.DataFrame(
            request_df['reg_doc_version_placebo'].apply(extract_root_and_latest).tolist(), 
            index=request_df.index
        )
        
        print("Active reg doc roots:", request_df['reg_doc_version_active_root'].unique().tolist())
        print("Active reg doc latest versions:", request_df['reg_doc_version_active_latest'].unique().tolist())
        print("Placebo reg doc roots:", request_df['reg_doc_version_placebo_root'].unique().tolist())
        print("Placebo reg doc latest versions:", request_df['reg_doc_version_placebo_latest'].unique().tolist())
        
        # Get Egnyte access token
        access_token = get_egnyte_token()
        if not access_token:
            return jsonify({
                "status": "error",
                "message": "Failed to get Egnyte access token due to rate limiting. Please try again later.",
                "error_code": "EGNYTE_RATE_LIMIT",
                "retry_suggestion": "Wait 10-15 minutes before retrying bulk requests"
            }), 429  # 429 = Too Many Requests
        
        # Get templates from Egnyte
        templates_folder_id = "966281ab-54c3-47ea-b20f-b38ed2ef9b30"
        logger.info(f"Fetching templates from folder ID: {templates_folder_id}")
        templates_data = list_egnyte_folder_contents(access_token, templates_folder_id)
        if not templates_data:
            logger.error("FAILED: Could not list templates folder contents")
            return jsonify({"error": "Failed to list templates folder contents"}), 500
        
        templates = templates_data.get("files", [])
        logger.info(f"Found {len(templates)} templates in Egnyte")
        for i, template in enumerate(templates[:10]):  # Log first 10 templates
            logger.info(f"Template {i+1}: {template.get('name')} (ID: {template.get('entry_id')})")
        
        # Get source documents from Egnyte
        source_docs_folder_id = "56545792-6b5d-4fc3-8e78-31d401bd7088"
        logger.info(f"Fetching source documents from folder ID: {source_docs_folder_id}")
        source_docs_data = list_egnyte_folder_contents(access_token, source_docs_folder_id)
        if not source_docs_data:
            logger.error("FAILED: Could not list source documents folder contents")
            return jsonify({"error": "Failed to list source documents folder contents"}), 500
        
        source_docs = source_docs_data.get("files", [])
        logger.info(f"Found {len(source_docs)} source documents in Egnyte")
        for i, doc in enumerate(source_docs[:10]):  # Log first 10 source docs
            logger.info(f"Source doc {i+1}: {doc.get('name')} (ID: {doc.get('entry_id')})")
        
        # Filter to only latest versions and process each request row
        latest_version_rows = []
        status_report = []
        
        for idx, row in request_df.iterrows():
            print(f"\n--- Processing Row {idx + 1} ---")
            print(f"Product Code: {row['product_code']}")
            print(f"Active Reg Doc Root: {row['reg_doc_version_active_root']}")
            print(f"Active Reg Doc Latest: {row['reg_doc_version_active_latest']}")
            print(f"Placebo Reg Doc Root: {row['reg_doc_version_placebo_root']}")
            print(f"Placebo Reg Doc Latest: {row['reg_doc_version_placebo_latest']}")
            print(f"Section: {row['section']}")
            
            # Check if this is the latest version
            active_is_latest = row['reg_doc_version_active_latest'] is not None
            placebo_is_latest = row['reg_doc_version_placebo_latest'] is not None
            
            if not active_is_latest and not placebo_is_latest:
                print(f"  ⚠️ Skipping - not latest version")
                status_report.append({
                    'row_index': idx,
                    'row_data': row.to_dict(),
                    'matching_source_document': None,
                    'matching_template': None,
                    'status': "old version"
                })
                continue
            
            # Find matching templates
            matching_templates = []
            logger.info(f"Looking for templates matching active_root: '{row['reg_doc_version_active_root']}' or placebo_root: '{row['reg_doc_version_placebo_root']}'")
            
            for template in templates:
                template_name = template.get('name', '').lower()
                active_root = row['reg_doc_version_active_root']
                placebo_root = row['reg_doc_version_placebo_root']
                
                # Convert to lowercase only if not None
                active_root_lower = active_root.lower() if active_root else ""
                placebo_root_lower = placebo_root.lower() if placebo_root else ""
                
                logger.info(f"Checking template: '{template.get('name')}' (ID: {template.get('entry_id')})")
                logger.info(f"  Template name (lower): '{template_name}'")
                logger.info(f"  Active root (lower): '{active_root_lower}'")
                logger.info(f"  Placebo root (lower): '{placebo_root_lower}'")
                
                # More specific template matching - check for exact section match
                template_matches = False
                if active_root_lower and active_root_lower in template_name:
                    template_matches = True
                    logger.info(f"  ✓ Template match (active): {template.get('name')} - matches {active_root_lower}")
                elif placebo_root_lower and placebo_root_lower in template_name:
                    template_matches = True
                    logger.info(f"  ✓ Template match (placebo): {template.get('name')} - matches {placebo_root_lower}")
                else:
                    logger.info(f"  ✗ No match for template: {template.get('name')}")
                
                if template_matches:
                    matching_templates.append(template)
            
            # Find matching source documents
            matching_source_docs = []
            logger.info(f"Looking for source documents matching product_code: '{row['product_code']}'")
            
            for source_doc in source_docs:
                source_doc_name = source_doc.get('name', '').lower()
                product_code = row['product_code'].lower()
                
                logger.info(f"Checking source doc: '{source_doc.get('name')}' (ID: {source_doc.get('entry_id')})")
                logger.info(f"  Source doc name (lower): '{source_doc_name}'")
                logger.info(f"  Product code (lower): '{product_code}'")
                
                # Check if source document matches the product code
                if product_code in source_doc_name:
                    matching_source_docs.append(source_doc)
                    logger.info(f"  ✓ Source doc match: {source_doc.get('name')} - matches product code {product_code}")
                else:
                    logger.info(f"  ✗ No match for source doc: {source_doc.get('name')}")
            
            # Determine status and add to report
            matching_template = matching_templates[0] if matching_templates else None
            matching_source_doc = matching_source_docs[0] if matching_source_docs else None
            
            # Debug the matched files
            if matching_template:
                logger.info(f"MATCHED TEMPLATE DETAILS:")
                logger.info(f"  Name: {matching_template.get('name')}")
                logger.info(f"  Entry ID: {matching_template.get('entry_id')}")
                logger.info(f"  Path: {matching_template.get('path')}")
                logger.info(f"  Type: {matching_template.get('type')}")
                logger.info(f"  Size: {matching_template.get('size')}")
            
            if matching_source_doc:
                logger.info(f"MATCHED SOURCE DOC DETAILS:")
                logger.info(f"  Name: {matching_source_doc.get('name')}")
                logger.info(f"  Entry ID: {matching_source_doc.get('entry_id')}")
                logger.info(f"  Path: {matching_source_doc.get('path')}")
                logger.info(f"  Type: {matching_source_doc.get('type')}")
                logger.info(f"  Size: {matching_source_doc.get('size')}")
            
            if matching_template and matching_source_doc:
                status = "Matched Both Docs in Egnyte"
            elif not matching_source_doc:
                status = f"Did not match source for product code {row['product_code']}"
            elif not matching_template:
                doc_code = row['reg_doc_version_active_root'] or row['reg_doc_version_placebo_root']
                status = f"did not match template {doc_code}"
            else:
                status = f"did not match to either source {row['product_code']} or template"
            
            status_report.append({
                'row_index': idx,
                'row_data': row.to_dict(),
                'matching_source_document': matching_source_doc,
                'matching_template': matching_template,
                'status': status
            })
            
            # Add to latest version rows if it has matches
            if matching_template or matching_source_doc:
                latest_version_rows.append({
                    'row_index': idx,
                    'row_data': row.to_dict(),
                    'matching_templates': matching_templates,
                    'matching_source_docs': matching_source_docs,
                    'total_templates': len(matching_templates),
                    'total_source_docs': len(matching_source_docs)
                })
            
            print(f"  Row {idx + 1} Summary: {len(matching_templates)} templates, {len(matching_source_docs)} source docs")
            print(f"  Status: {status}")
        
        # Filter status report to only show "Matched Both Docs in Egnyte"
        matched_status_report = [item for item in status_report if item['status'] == "Matched Both Docs in Egnyte"]
        
        # Create summary table of status reasons
        status_counts = {}
        for item in status_report:
            status = item['status']
            status_counts[status] = status_counts.get(status, 0) + 1
        
        summary_table = []
        for status, count in status_counts.items():
            summary_table.append({
                'status': status,
                'count': count,
                'percentage': round((count / len(status_report)) * 100, 1)
            })
        
        # Print overall summary
        total_latest_matches = len(latest_version_rows)
        total_matches = sum(len(match['matching_templates']) + len(match['matching_source_docs']) for match in latest_version_rows)
        print(f"\n=== OVERALL SUMMARY ===")
        print(f"Total rows processed: {len(request_df)}")
        print(f"Latest version rows with matches: {total_latest_matches}")
        print(f"Total matches found: {total_matches}")
        print(f"Rows with 'Matched Both Docs in Egnyte': {len(matched_status_report)}")
        
        print(f"\n=== STATUS SUMMARY TABLE ===")
        for summary in summary_table:
            print(f"{summary['status']}: {summary['count']} rows ({summary['percentage']}%)")
        
        # Process document generation for matched rows
        document_generation_results = []
        generated_docs_urls = []
        
        for matched_row in matched_status_report:
            logger.info("=" * 80)
            logger.info("PROCESSING MATCHED ROW FOR DOCUMENT GENERATION")
            logger.info("=" * 80)
            logger.info(f"Row index: {matched_row['row_index']}")
            logger.info(f"Product code: {matched_row['row_data']['product_code']}")
            logger.info(f"Status: {matched_row['status']}")
            
            if matched_row['matching_template']:
                logger.info(f"Template to use: {matched_row['matching_template'].get('name')} (ID: {matched_row['matching_template'].get('entry_id')})")
            else:
                logger.info("No template found!")
                
            if matched_row['matching_source_document']:
                logger.info(f"Source doc to use: {matched_row['matching_source_document'].get('name')} (ID: {matched_row['matching_source_document'].get('entry_id')})")
            else:
                logger.info("No source document found!")
            
            generation_result = process_document_generation(matched_row)
            
            # Extract URLs if generation was successful
            doc_urls = {}
            if generation_result.get('success'):
                upload_result = generation_result.get('upload_result', {})
                docx_result = upload_result.get('docx_result', {})
                pdf_result = upload_result.get('pdf_result', {})
                
                # Build full URLs
                if docx_result:
                    docx_url = f"https://{DOMAIN}/app/index.do#storage/files/1{docx_result.get('path', '')}"
                    doc_urls['docx_url'] = docx_url
                
                if pdf_result:
                    pdf_url = f"https://{DOMAIN}/app/index.do#storage/files/1{pdf_result.get('path', '')}"
                    doc_urls['pdf_url'] = pdf_url
            
            document_generation_results.append({
                'row_index': matched_row['row_index'],
                'product_code': matched_row['row_data']['product_code'],
                'generation_result': generation_result,
                'doc_urls': doc_urls
            })
            
            # Add to generated docs URLs list
            if doc_urls:
                generated_docs_urls.append({
                    'product_code': matched_row['row_data']['product_code'],
                    'section': matched_row['row_data']['section'],
                    'docx_filename': generation_result.get('docx_filename'),
                    'pdf_filename': generation_result.get('pdf_filename'),
                    'docx_url': doc_urls.get('docx_url'),
                    'pdf_url': doc_urls.get('pdf_url')
                })
        
        # Create total match report
        total_match_report = {
            'campaign_summary': {
                'total_requests': len(request_df),
                'successful_matches': len(matched_status_report),
                'unique_product_codes': unique_product_codes,
                'processing_timestamp': timestamp_clean
            },
            'status_breakdown': summary_table,
            'generated_documents': generated_docs_urls,
            'detailed_results': document_generation_results
        }
        
        return jsonify({
            "status": "success", 
            "message": "Bulk request processed successfully",
            "total_match_report": total_match_report
        }), 200
        
    except Exception as e:
        logger.error(f"Error in reg_docs_bulk_request: {e}")
        return jsonify({"error": str(e)}), 500

# Modular Functions for Document Processing Workflow

def load_prompt_from_file():
    """Load the prompt from demo_prompt.py"""
    try:
        # Try UTF-8 first, then fall back to other encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open('demo_prompt.py', 'r', encoding=encoding) as file:
                    content = file.read()
                    # Extract the prompt variable content
                    start = content.find('prompt = """') + 11
                    end = content.find('"""', start)
                    if start > 10 and end > start:  # Valid indices
                        prompt = content[start:end]
                        logger.info(f"Successfully loaded prompt with {encoding} encoding ({len(prompt)} characters)")
                        return prompt
                    else:
                        logger.warning(f"Could not find prompt in file with {encoding} encoding")
            except UnicodeDecodeError:
                logger.warning(f"Failed to decode with {encoding} encoding, trying next...")
                continue
        
        # If all encodings fail, try with error handling
        try:
            with open('demo_prompt.py', 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                start = content.find('prompt = """') + 11
                end = content.find('"""', start)
                if start > 10 and end > start:
                    prompt = content[start:end]
                    logger.info(f"Loaded prompt with error handling ({len(prompt)} characters)")
                    return prompt
        except Exception as e:
            logger.error(f"Failed to load prompt even with error handling: {e}")
        
        # Fallback prompt if file loading fails
        logger.warning("Using fallback prompt due to encoding issues")
        fallback_prompt = """AI-Prompt - IND Module 3: Section 3.2.P.1 Draft

Context
You are a regulatory-writing assistant drafting Section 3.2.P.1 "Description and Composition of the Drug Product" for an IND that is currently in Phase II clinical trials. All content must be suitable for direct inclusion in an eCTD-compliant Module 3 dossier.

Product Information Provided Separately
- Product code: <PRODUCT_CODE>
- Dosage form: Immediate-release film-coated tablet
- Active strength(s) and qualitative/quantitative composition appear in the source data you have received.

Required Output Structure & Style

1. 3.2.P.1.1 Description of the Dosage Form

- One concise paragraph that:
  - Identifies the dosage form and strength(s)
  - States the active-ingredient concentration(s) clearly.
  - Summarises the mechanism of action in one sentence.

- Write in the third person, scientific style; do not use marketing language.

2. 3.2.P.1.2 Composition

- Introductory sentence: "The qualitative and quantitative composition of the is provided in Table 1."

- Table 1 - 'Composition of the '

  Columns (exact wording):
  1. Component
  2. Quality Reference
  3. Function
  4. Quantity / Unit (mg per tablet)

- Populate all excipients and active ingredient(s) as rows using the supplied source data.
- Add a Total Weight row.
- Place a footnote directly beneath the table:
  "Abbreviations: NF = National Formulary; Ph. Eur. = European Pharmacopoeia; USP = United States Pharmacopoeia."

Formatting Instructions
- Use standard IND section numbering and bold headings exactly as shown.
- Table should render in Word/PDF without manual re-formatting (plain grid, no shading).
- Do not add any content outside Sections 3.2.P.1.1 and 3.2.P.1.2.

Quality Criteria
- All numeric values must match the source composition data.
- Spelling per US English; pharmacopoeia references capitalised.
- Ready for QC with minimal edits.

Please draft the section now, substituting the actual composition data in Table 1."""
        
        return fallback_prompt
        
    except Exception as e:
        logger.error(f"Error loading prompt: {e}")
        return None

def download_egnyte_file_to_temp(access_token, file_id, file_extension='.tmp', file_path=None):
    """Download a file from Egnyte to a temporary file"""
    try:
        logger.info(f"Attempting to download file {file_id} with extension {file_extension}")
        if file_path:
            logger.info(f"File path available: {file_path}")
        
        # Download file content
        file_content = download_egnyte_file(access_token, file_id, file_path)
        if not file_content:
            logger.error(f"download_egnyte_file returned None for file_id: {file_id}")
            return None
        
        logger.info(f"Successfully downloaded {len(file_content)} bytes for file_id: {file_id}")
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
        temp_file.write(file_content)
        temp_file.close()
        
        logger.info(f"Created temporary file: {temp_file.name}")
        return temp_file.name
    except Exception as e:
        logger.error(f"Error downloading file to temp: {e}")
        return None

def upload_files_prompt_to_openai(prompt: str, template_path: str, source_document_path: str) -> str:
    """
    Upload files to OpenAI and generate a document using the prompt and uploaded files.
    
    Args:
        prompt: The prompt text to guide document generation
        template_path: Path to the template file (DOCX or PDF)
        source_document_path: Path to the source document file (DOCX or PDF)
    
    Returns:
        Generated document content in DOCX format as bytes
    """
    try:
        logger.info("=" * 60)
        logger.info("STARTING DOCUMENT GENERATION WITH OPENAI")
        logger.info("=" * 60)
        logger.info(f"Prompt length: {len(prompt)} characters")
        logger.info(f"Template path: {template_path}")
        logger.info(f"Source document path: {source_document_path}")
        
        # Initialize OpenAI client
        client = initialize_openai()
        if not client:
            logger.error("FAILED: Could not initialize OpenAI client")
            return None
        logger.info("SUCCESS: OpenAI client initialized")
        
        # Convert DOCX to PDF if needed for template (Responses API requires PDF)
        template_file_path = template_path
        if template_path.lower().endswith('.docx'):
            logger.info("Converting template DOCX to PDF for upload...")
            template_pdf_path = convert_docx_to_pdf_for_upload(template_path)
            if template_pdf_path:
                template_file_path = template_pdf_path
                logger.info(f"Template converted to PDF: {template_pdf_path}")
            else:
                logger.error("Failed to convert template to PDF")
                return None
        
        # Convert DOCX to PDF if needed for source document (Responses API requires PDF)
        source_file_path = source_document_path
        if source_document_path.lower().endswith('.docx'):
            logger.info("Converting source document DOCX to PDF for upload...")
            source_pdf_path = convert_docx_to_pdf_for_upload(source_document_path)
            if source_pdf_path:
                source_file_path = source_pdf_path
                logger.info(f"Source document converted to PDF: {source_pdf_path}")
            else:
                logger.error("Failed to convert source document to PDF")
                return None
        
        # Upload template file
        logger.info("Uploading template file to OpenAI...")
        with open(template_file_path, 'rb') as file:
            template_file = client.files.create(
                file=file,
                purpose='user_data'
            )
            template_file_id = template_file.id
        logger.info(f"SUCCESS: Template file uploaded with ID: {template_file_id}")
        
        # Upload source document file
        logger.info("Uploading source document file to OpenAI...")
        with open(source_file_path, 'rb') as file:
            source_file = client.files.create(
                file=file,
                purpose='user_data'
            )
            source_file_id = source_file.id
        logger.info(f"SUCCESS: Source document file uploaded with ID: {source_file_id}")
        
        # Generate document using OpenAI Responses API
        logger.info("Generating document with OpenAI...")
        response = client.responses.create(
            model="gpt-4o",
            input=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_file",
                            "file_id": template_file_id
                        },
                        {
                            "type": "input_file",
                            "file_id": source_file_id
                        },
                        {
                            "type": "input_text",
                            "text": f"""
{prompt}

INSTRUCTIONS:
1. Use the template file as the structure and format for the new document
2. Extract relevant information from the source document
3. Create a comprehensive, well-structured document following the template format
4. Maintain professional tone and regulatory compliance
5. Return the document in clean, properly formatted HTML that can be converted to DOCX

FORMATTING REQUIREMENTS:
- Use <h1> for main headings, <h2> for subheadings, <h3> for sub-subheadings
- Use <ul><li> for bullet points and lists
- Use <strong> for important terms, specifications, and key data
- Use <p> for paragraphs
- Use <table><tr><td> for tables with proper structure
- Use clear, professional language without placeholders or brackets
- Structure information logically with proper spacing
- Include all relevant pharmaceutical data from the source document
- Make the document ready for immediate regulatory use
- Return complete, valid HTML that can be directly rendered

Please generate the complete document content in HTML format based on the template and source document.
"""
                        }
                    ]
                }
            ]
        )
        
        # Extract the generated content
        generated_content = response.output_text
        logger.info(f"SUCCESS: Document generated ({len(generated_content)} characters)")
        
        # Convert the generated HTML content to DOCX format
        logger.info("Converting generated HTML content to DOCX format...")
        
        # Create a temporary file path for the DOCX
        import tempfile
        temp_docx_path = tempfile.mktemp(suffix='.docx')
        
        # Convert HTML to DOCX using the proper HTML converter
        success = convert_html_to_docx(generated_content, temp_docx_path)
        if not success:
            logger.error("Failed to convert generated HTML content to DOCX")
            return None
        
        # Read the generated DOCX file
        with open(temp_docx_path, 'rb') as docx_file:
            docx_content = docx_file.read()
        
        # Clean up temporary file
        try:
            os.unlink(temp_docx_path)
        except:
            pass
        
        # Clean up uploaded files
        logger.info("Cleaning up uploaded files...")
        try:
            client.files.delete(template_file_id)
            client.files.delete(source_file_id)
            logger.info("SUCCESS: Uploaded files cleaned up")
        except Exception as e:
            logger.warning(f"Warning: Could not delete uploaded files: {e}")
        
        # Clean up temporary PDF files if they were created
        if template_file_path != template_path and os.path.exists(template_file_path):
            os.unlink(template_file_path)
        if source_file_path != source_document_path and os.path.exists(source_file_path):
            os.unlink(source_file_path)
        
        logger.info("SUCCESS: Document generation completed")
        logger.info("=" * 60)
        
        return docx_content
        
    except Exception as e:
        logger.error("=" * 60)
        logger.error("DOCUMENT GENERATION FAILED")
        logger.error("=" * 60)
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
        logger.error("=" * 60)
        return None

def convert_docx_to_pdf_for_upload(docx_path: str) -> str:
    """Convert DOCX file to PDF with professional formatting"""
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        # Extract content from DOCX with structure preservation
        doc = Document(docx_path)
        
        # Create PDF with professional settings
        pdf_path = docx_path.replace('.docx', '_temp.pdf')
        doc_pdf = SimpleDocTemplate(
            pdf_path, 
            pagesize=letter,
            topMargin=1*inch,
            bottomMargin=1*inch,
            leftMargin=1*inch,
            rightMargin=1*inch
        )
        
        # Define professional styles
        styles = getSampleStyleSheet()
        
        # Custom title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            alignment=1,  # Center alignment
            textColor=colors.HexColor('#1f497d')  # Dark blue
        )
        
        # Custom heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.HexColor('#1f497d')  # Dark blue
        )
        
        # Custom normal style
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        )
        
        story = []
        
        # Process document content - simplified approach
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    if level == 1:
                        story.append(Paragraph(text, title_style))
                    else:
                        story.append(Paragraph(text, heading_style))
                else:
                    story.append(Paragraph(text, normal_style))
        
        # Process tables separately
        for table in doc.tables:
            # Convert table to ReportLab format
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                # Create ReportLab table
                rl_table = Table(table_data)
                
                # Apply professional table styling
                table_style = TableStyle([
                    # Header row styling
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),  # Blue background
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # White text
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    
                    # Data rows styling
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F2F2F2')),  # Light gray
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    
                    # Grid and alignment
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F8F8')]),
                ])
                
                rl_table.setStyle(table_style)
                story.append(rl_table)
                story.append(Spacer(1, 12))  # Add spacing after table
        
        # Build PDF
        doc_pdf.build(story)
        logger.info(f"SUCCESS: PDF created with professional formatting: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {e}")
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
        return None

def convert_text_to_docx(text_content: str) -> bytes:
    """Convert text content to DOCX format with proper formatting"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import io
        import re
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('Regulatory Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content line by line for better formatting
        lines = text_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Handle different types of content
            if line.startswith('# '):
                # Main heading
                heading_text = line[2:].strip()
                doc.add_heading(heading_text, 1)
                
            elif line.startswith('## '):
                # Sub heading
                heading_text = line[3:].strip()
                doc.add_heading(heading_text, 2)
                
            elif line.startswith('### '):
                # Sub-sub heading
                heading_text = line[4:].strip()
                doc.add_heading(heading_text, 3)
                
            elif line.startswith('- **') and line.endswith('**'):
                # Bold list item
                item_text = line[2:].strip()  # Remove "- "
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)  # Remove ** markers
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(item_text).bold = True
                
            elif line.startswith('- '):
                # Regular list item
                item_text = line[2:].strip()
                p = doc.add_paragraph()
                p.add_run('• ')
                p.add_run(item_text)
                
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                bold_text = line[2:-2].strip()
                p = doc.add_paragraph()
                p.add_run(bold_text).bold = True
                
            elif line.startswith('Table'):
                # Table reference
                p = doc.add_paragraph()
                p.add_run(line).italic = True
                
            elif ':' in line and len(line.split(':')) == 2:
                # Key-value pairs
                key, value = line.split(':', 1)
                p = doc.add_paragraph()
                p.add_run(key.strip() + ': ').bold = True
                p.add_run(value.strip())
                
            else:
                # Regular paragraph
                # Clean up any remaining markdown
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Remove ** markers
                clean_text = re.sub(r'\[.*?\]', '', clean_text)  # Remove [placeholder] text
                clean_text = clean_text.strip()
                
                if clean_text and clean_text != line.strip():
                    # If we cleaned up markdown, make it bold
                    p = doc.add_paragraph()
                    p.add_run(clean_text).bold = True
                elif clean_text:
                    # Regular paragraph
                    doc.add_paragraph(clean_text)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
        
    except Exception as e:
        logger.error(f"Error converting text to DOCX: {e}")
        return None

def extract_row_data(row_data):
    """Extract relevant data from the row"""
    try:
        relevant_data = {
            'ctm': row_data.get('ctm'),
            'product_code': row_data.get('product_code'),
            'mfg_lot': row_data.get('mfg_lot'),
            'mfg_type': row_data.get('mfg_type'),
            'section': row_data.get('section'),
            'filing_type': row_data.get('filing_type'),
            'project': row_data.get('project'),
            'module': row_data.get('module')
        }
        return relevant_data
    except Exception as e:
        logger.error(f"Error extracting row data: {e}")
        return None

def convert_html_to_docx(html_content, output_path):
    """Convert HTML content to DOCX format with professional formatting"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        from bs4 import BeautifulSoup
        import re
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create Word document with professional settings
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Define professional styles
        def add_heading_with_style(doc, text, level):
            """Add heading with professional styling"""
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 73, 125)  # Dark blue
            return heading
        
        def add_paragraph_with_style(doc, text, bold=False, italic=False):
            """Add paragraph with professional styling"""
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = bold
            run.font.italic = italic
            return p
        
        # Process HTML content with better structure handling
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'div']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                text = element.get_text().strip()
                if text:
                    add_heading_with_style(doc, text, level)
                    
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    # Check for bold or italic formatting
                    is_bold = bool(element.find(['strong', 'b']))
                    is_italic = bool(element.find(['em', 'i']))
                    add_paragraph_with_style(doc, text, bold=is_bold, italic=is_italic)
                    
            elif element.name == 'table':
                # Enhanced table handling
                rows = element.find_all('tr')
                if rows:
                    # Determine table dimensions
                    max_cols = 0
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        max_cols = max(max_cols, len(cells))
                    
                    if max_cols > 0:
                        # Create table with proper dimensions
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        # Apply professional table styling
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if i < len(table.rows) and j < len(table.rows[i].cells):
                                    cell_text = cell.get_text().strip()
                                    table_cell = table.rows[i].cells[j]
                                    table_cell.text = cell_text
                                    
                                    # Style header row
                                    if i == 0 or cell.name == 'th':
                                        for paragraph in table_cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.font.bold = True
                                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                                run.font.size = Pt(10)
                                        # Set header background color
                                        table_cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                                        table_cell._tc.get_or_add_tcPr().xpath('w:shd')[0].set(qn('w:fill'), '4472C4')  # Blue background
                                    else:
                                        # Style data rows
                                        for paragraph in table_cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.font.size = Pt(9)
                                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                        
                        # Add spacing after table
                        doc.add_paragraph()
                        
            elif element.name == 'div':
                # Handle div elements that might contain structured content
                text = element.get_text().strip()
                if text and not element.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table']):
                    add_paragraph_with_style(doc, text)
        
        # Add professional footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph("Document generated by Regulatory Document System")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        
        # Save document
        doc.save(output_path)
        logger.info(f"SUCCESS: DOCX file created with professional formatting: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error converting HTML to DOCX: {e}")
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
        return False

def upload_generated_files_to_egnyte(access_token, docx_path, pdf_path, folder_id):
    """Upload generated files to Egnyte"""
    try:
        # Upload DOCX file
        with open(docx_path, 'rb') as docx_file:
            docx_content = docx_file.read()
            docx_result = upload_file_to_egnyte(access_token, folder_id, os.path.basename(docx_path), docx_content)
        
        # Upload PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()
            pdf_result = upload_file_to_egnyte(access_token, folder_id, os.path.basename(pdf_path), pdf_content)
        
        return {
            'docx_uploaded': docx_result is not None,
            'pdf_uploaded': pdf_result is not None,
            'docx_result': docx_result,
            'pdf_result': pdf_result
        }
        
    except Exception as e:
        logger.error(f"Error uploading files to Egnyte: {e}")
        return None

def process_document_generation(matched_row):
    """Main function to process document generation for a matched row"""
    try:
        logger.info("=" * 80)
        logger.info("STARTING DOCUMENT GENERATION PROCESS")
        logger.info("=" * 80)
        logger.info(f"Processing row: {matched_row.get('row_index')}")
        logger.info(f"Product code: {matched_row.get('row_data', {}).get('product_code')}")
        logger.info(f"Section: {matched_row.get('row_data', {}).get('section')}")
        
        # Get Egnyte access token
        logger.info("Step 1: Getting Egnyte access token...")
        access_token = get_egnyte_token()
        if not access_token:
            logger.error("FAILED: Could not get Egnyte access token due to rate limiting")
            return {"error": "Failed to get Egnyte access token - rate limit exceeded. Please try again in 10-15 minutes."}
        logger.info("SUCCESS: Egnyte access token obtained")
        
        # Step 2: Load prompt
        logger.info("Step 2: Loading prompt from demo_prompt.py...")
        prompt = load_prompt_from_file()
        if not prompt:
            logger.error("FAILED: Could not load prompt from demo_prompt.py")
            return {"error": "Failed to load prompt"}
        logger.info(f"SUCCESS: Prompt loaded ({len(prompt)} characters)")
        
        # Step 3: Download template and source documents
        logger.info("Step 3: Downloading template and source documents...")
        template_file = matched_row['matching_template']
        source_file = matched_row['matching_source_document']
        
        if not template_file:
            logger.error("FAILED: No template file found in matched_row")
            return {"error": "No template file found"}
        
        if not source_file:
            logger.error("FAILED: No source document found in matched_row")
            return {"error": "No source document found"}
        
        logger.info(f"Template file: {template_file.get('name')} (ID: {template_file.get('entry_id')})")
        logger.info(f"Source file: {source_file.get('name')} (ID: {source_file.get('entry_id')})")
        
        # Download template file
        template_temp_path = download_egnyte_file_to_temp(access_token, template_file['entry_id'], '.docx', template_file.get('path'))
        if not template_temp_path:
            logger.error("FAILED: Could not download template file to temp location")
            return {"error": "Failed to download template file"}
        logger.info(f"SUCCESS: Template downloaded to {template_temp_path}")
        
        # Download source document
        source_temp_path = download_egnyte_file_to_temp(access_token, source_file['entry_id'], '.pdf', source_file.get('path'))
        if not source_temp_path:
            logger.error("FAILED: Could not download source document to temp location")
            return {"error": "Failed to download source document"}
        logger.info(f"SUCCESS: Source document downloaded to {source_temp_path}")
        
        # Step 4: Generate document using the new upload_files_prompt_to_openai function
        logger.info("Step 4: Generating document with OpenAI file upload...")
        docx_content = upload_files_prompt_to_openai(
            prompt=prompt,
            template_path=template_temp_path,
            source_document_path=source_temp_path
        )
        
        # Clean up downloaded temp files
        logger.info("Cleaning up downloaded temp files...")
        os.unlink(template_temp_path)
        os.unlink(source_temp_path)
        logger.info("Downloaded temp files cleaned up")
        
        if not docx_content:
            logger.error("FAILED: Could not generate document with OpenAI")
            return {"error": "Failed to generate document"}
        logger.info(f"SUCCESS: Document generated ({len(docx_content)} bytes)")
        
        # Step 5: Create output files
        logger.info("Step 5: Creating output files...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        product_code = matched_row['row_data']['product_code']
        docx_filename = f"{product_code}_regulatory_doc_{timestamp}.docx"
        pdf_filename = f"{product_code}_regulatory_doc_{timestamp}.pdf"
        
        docx_path = os.path.join(tempfile.gettempdir(), docx_filename)
        pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
        
        logger.info(f"DOCX path: {docx_path}")
        logger.info(f"PDF path: {pdf_path}")
        
        # Save DOCX content to file
        with open(docx_path, 'wb') as f:
            f.write(docx_content)
        logger.info("SUCCESS: DOCX file created")
        
        # Convert DOCX to PDF
        logger.info("Converting DOCX to PDF...")
        pdf_success = convert_docx_to_pdf_for_upload(docx_path)
        if pdf_success and os.path.exists(pdf_success):
            # Move the PDF to the expected location
            import shutil
            shutil.move(pdf_success, pdf_path)
            logger.info("SUCCESS: PDF file created")
        else:
            logger.error("FAILED: Could not convert DOCX to PDF")
            return {"error": "Failed to convert to PDF"}
        
        # Step 6: Upload to Egnyte
        logger.info("Step 6: Uploading files to Egnyte...")
        
        # Find the target folder dynamically
        molecule_code = matched_row['row_data'].get('molecule_code', 'THPG001')
        campaign_number = matched_row['row_data'].get('campaign_number', '4')
        
        logger.info(f"Looking for target folder for molecule: {molecule_code}, campaign: {campaign_number}")
        target_folder_id = find_egnyte_target_folder(access_token, molecule_code, campaign_number)
        
        if not target_folder_id:
            logger.warning("Could not find target folder, using fallback folder")
            # Fallback to a known folder - you might want to update this
            target_folder_id = "4a85f5e6-bb31-4bd1-b011-6fc75bdcb2d7"
        
        logger.info(f"Target folder ID: {target_folder_id}")
        
        upload_result = upload_generated_files_to_egnyte(access_token, docx_path, pdf_path, target_folder_id)
        
        # Clean up temp files
        logger.info("Cleaning up temporary files...")
        os.unlink(docx_path)
        os.unlink(pdf_path)
        logger.info("Temporary files cleaned up")
        
        if not upload_result:
            logger.error("FAILED: Could not upload files to Egnyte")
            return {"error": "Failed to upload files to Egnyte"}
        
        logger.info(f"SUCCESS: Files uploaded to Egnyte - {upload_result}")
        logger.info("=" * 80)
        logger.info("DOCUMENT GENERATION PROCESS COMPLETED SUCCESSFULLY")
        logger.info("=" * 80)
        
        return {
            "success": True,
            "docx_filename": docx_filename,
            "pdf_filename": pdf_filename,
            "upload_result": upload_result
        }
        
    except Exception as e:
        logger.error("=" * 80)
        logger.error("DOCUMENT GENERATION PROCESS FAILED")
        logger.error("=" * 80)
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        logger.error(f"Error details: {e}")
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
        logger.error("=" * 80)
        return {"error": str(e)}


if __name__ == '__main__':
    # For local development
    app.run(debug=False, host='0.0.0.0', port=5000) 

