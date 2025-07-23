import streamlit as st
import pandas as pd
import io
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import tempfile
import os
from openai import OpenAI
from typing import Dict, List, Optional
import plotly.express as px
import plotly.graph_objects as go
from PIL import Image
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.io as pio
import numpy as np
import json
from google.oauth2 import service_account
from googleapiclient.discovery import build

# Page configuration
st.set_page_config(
    page_title="Document Generator and Folder Management System",
    page_icon="📁",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: #2c3e50;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .stButton > button {
        width: 100%;
        margin-top: 1rem;
    }
    .regulatory-text {
        background-color: #f8f9fa;
        padding: 1rem;
        border-left: 4px solid #1f77b4;
        margin: 1rem 0;
    }
    .chart-container {
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .drug-info {
        background-color: #e8f4fd;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Drug database
DRUG_DATABASE = {
    "Metformin": {
        "dosage_form": "Immediate-release film-coated tablet",
        "mechanism": "Metformin is a biguanide antihyperglycemic agent that improves glucose tolerance in patients with type 2 diabetes, lowering both basal and postprandial plasma glucose. Its pharmacologic mechanisms of action are different from other classes of oral antihyperglycemic agents.",
        "indication": "Type 2 diabetes mellitus",
        "strength": "500 mg, 850 mg, 1000 mg",
        "class": "Biguanide",
        "manufacturer": "Various"
    },
    "Atorvastatin": {
        "dosage_form": "Film-coated tablet",
        "mechanism": "Atorvastatin is a selective, competitive inhibitor of HMG-CoA reductase, the rate-limiting enzyme that converts 3-hydroxy-3-methylglutaryl-coenzyme A to mevalonate, a precursor of sterols, including cholesterol.",
        "indication": "Hypercholesterolemia and cardiovascular risk reduction",
        "strength": "10 mg, 20 mg, 40 mg, 80 mg",
        "class": "HMG-CoA reductase inhibitor (statin)",
        "manufacturer": "Various"
    },
    "Lisinopril": {
        "dosage_form": "Film-coated tablet",
        "mechanism": "Lisinopril is a competitive inhibitor of angiotensin-converting enzyme (ACE). ACE is a peptidyl dipeptidase that catalyzes the conversion of angiotensin I to the vasoconstrictor substance, angiotensin II.",
        "indication": "Hypertension, heart failure, myocardial infarction",
        "strength": "2.5 mg, 5 mg, 10 mg, 20 mg, 30 mg, 40 mg",
        "class": "ACE inhibitor",
        "manufacturer": "Various"
    },
    "Omeprazole": {
        "dosage_form": "Delayed-release capsule",
        "mechanism": "Omeprazole is a proton pump inhibitor that suppresses gastric acid secretion by specific inhibition of the H+/K+-ATPase in the gastric parietal cell.",
        "indication": "Gastroesophageal reflux disease, peptic ulcer disease",
        "strength": "10 mg, 20 mg, 40 mg",
        "class": "Proton pump inhibitor",
        "manufacturer": "Various"
    },
    "Custom Drug": {
        "dosage_form": "Immediate-release film-coated tablet",
        "mechanism": "The active ingredient selectively inhibits the target enzyme, leading to therapeutic effects in the treatment of the indicated condition.",
        "indication": "Custom indication",
        "strength": "25 mg",
        "class": "Custom class",
        "manufacturer": "Custom manufacturer"
    }
}

def load_openai_api_key():
    """Load OpenAI API key from credentials file"""
    try:
        # Try local credentials file
        import credentials
        return credentials.OPENAI_API_KEY
    except:
        return None

def load_google_drive_credentials():
    """Load Google Drive API credentials from local JSON file or Streamlit secrets"""
    try:
        # First, try to load from Streamlit secrets
        if hasattr(st.secrets, 'google_drive_api'):
            credentials_dict = st.secrets.google_drive_api
            creds = service_account.Credentials.from_service_account_info(
                credentials_dict, 
                scopes=['https://www.googleapis.com/auth/drive']
            )
            return creds
        
        # If not in secrets, try to load from local JSON file
        json_file_path = 'aaitdemoharmony-3945571299f1.json'
        if os.path.exists(json_file_path):
            creds = service_account.Credentials.from_service_account_file(
                json_file_path, 
                scopes=['https://www.googleapis.com/auth/drive']
            )
            return creds
        
        # If neither exists, return None
        return None
        
    except Exception as e:
        st.error(f"Error loading Google Drive credentials: {e}")
        return None

def get_shared_drive_id(service):
    """Get the first available shared drive ID"""
    try:
        # List shared drives
        drives = service.drives().list(pageSize=10).execute()
        shared_drives = drives.get('drives', [])
        
        if shared_drives:
            # Return the first shared drive ID
            return shared_drives[0]['id']
        else:
            st.warning("⚠️ No shared drives found. File uploads may fail due to service account storage limitations.")
            st.info("💡 To fix this: Create a shared drive and add the service account as a member with 'Editor' permissions")
            return None
    except Exception as e:
        st.warning(f"⚠️ Error accessing shared drives: {e}")
        st.info("💡 This may be due to permissions. Try creating a shared drive manually.")
        return None

def initialize_google_drive_service():
    """Initialize Google Drive service with credentials"""
    creds = load_google_drive_credentials()
    if creds:
        try:
            service = build('drive', 'v3', credentials=creds)
            return service
        except Exception as e:
            st.error(f"Error initializing Google Drive service: {e}")
            return None
    return None

def list_google_drive_folders(service, parent_folder_id: str = None):
    """List folders in Google Drive, optionally within a specific parent folder"""
    try:
        if parent_folder_id:
            query = f"'{parent_folder_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        else:
            query = "mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        
        results = service.files().list(q=query, fields="files(id, name, createdTime, trashed)").execute()
        folders = results.get('files', [])
        
        # Filter out trashed folders
        active_folders = [folder for folder in folders if not folder.get('trashed', False)]
        return active_folders
    except Exception as e:
        st.error(f"Error listing folders: {e}")
        return []

def create_google_drive_folder(service, folder_name: str, parent_folder_id: str = None, shared_drive_id: str = None):
    """Create a new folder in Google Drive (shared drive or personal)"""
    try:
        folder_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        if parent_folder_id:
            folder_metadata['parents'] = [parent_folder_id]
        
        # If using shared drive, specify the drive ID
        if shared_drive_id:
            folder = service.files().create(
                body=folder_metadata, 
                fields='id, name',
                supportsAllDrives=True,
                supportsTeamDrives=True
            ).execute()
        else:
            folder = service.files().create(body=folder_metadata, fields='id, name').execute()
        
        return folder
    except Exception as e:
        st.error(f"Error creating folder '{folder_name}': {e}")
        return None

def get_folder_structure(service, parent_folder_id: str = None, max_depth: int = 3):
    """Get folder structure up to 3 levels deep"""
    try:
        if parent_folder_id:
            query = f"'{parent_folder_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        else:
            query = "mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        
        results = service.files().list(q=query, fields="files(id, name, parents, trashed)").execute()
        folders = results.get('files', [])
        
        # Filter out folders that are trashed
        active_folders = [folder for folder in folders if not folder.get('trashed', False)]
        
        structure = []
        for folder in active_folders:
            folder_info = {
                'id': folder['id'],
                'name': folder['name'],
                'level': 1,
                'children': []
            }
            
            # Get second level
            if max_depth > 1:
                level2_query = f"'{folder['id']}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
                level2_results = service.files().list(q=level2_query, fields="files(id, name, parents, trashed)").execute()
                level2_folders = level2_results.get('files', [])
                
                # Filter level 2 folders
                active_level2_folders = [f for f in level2_folders if not f.get('trashed', False)]
                
                for level2_folder in active_level2_folders:
                    level2_info = {
                        'id': level2_folder['id'],
                        'name': level2_folder['name'],
                        'level': 2,
                        'children': []
                    }
                    
                    # Get third level
                    if max_depth > 2:
                        level3_query = f"'{level2_folder['id']}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
                        level3_results = service.files().list(q=level3_query, fields="files(id, name, parents, trashed)").execute()
                        level3_folders = level3_results.get('files', [])
                        
                        # Filter level 3 folders
                        active_level3_folders = [f for f in level3_folders if not f.get('trashed', False)]
                        
                        for level3_folder in active_level3_folders:
                            level3_info = {
                                'id': level3_folder['id'],
                                'name': level3_folder['name'],
                                'level': 3,
                                'children': []
                            }
                            level2_info['children'].append(level3_info)
                    
                    folder_info['children'].append(level2_info)
            
            structure.append(folder_info)
        
        return structure
    except Exception as e:
        st.error(f"Error getting folder structure: {e}")
        return []

def get_folder_structure_recursive(service, parent_folder_id: str = None, max_depth: int = 3, current_depth: int = 0):
    """Get folder structure recursively to ensure proper hierarchy"""
    try:
        if current_depth >= max_depth:
            return []
        
        if parent_folder_id:
            query = f"'{parent_folder_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        else:
            query = f"'{parent_folder_id or '0ALsvNdCE73XrUk9PVA'}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        
        # For shared drives, we need to include the shared drive parameters
        # Always use shared drive parameters when we're in a shared drive context
        results = service.files().list(
            q=query, 
            fields="files(id, name, parents, trashed, createdTime, modifiedTime, owners, webViewLink)",
            supportsAllDrives=True,
            supportsTeamDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        folders = results.get('files', [])
        
        # Simplified filtering - just check if not trashed
        active_folders = []
        for folder in folders:
            # Check if folder is not trashed
            if folder.get('trashed', False):
                continue
            active_folders.append(folder)
        
        structure = []
        for folder in active_folders:
            folder_info = {
                'id': folder['id'],
                'name': folder['name'],
                'level': current_depth + 1,
                'webViewLink': folder.get('webViewLink', f"https://drive.google.com/drive/folders/{folder['id']}"),
                'children': []
            }
            
            # Recursively get children - always try to get them
            if current_depth < max_depth - 1:
                children = get_folder_structure_recursive(service, folder['id'], max_depth, current_depth + 1)
                folder_info['children'] = children
            
            structure.append(folder_info)
        
        return structure
    except Exception as e:
        st.error(f"Error getting folder structure: {e}")
        return []

def display_folder_structure(structure, level=0):
    """Display folder structure with proper indentation and clickable links"""
    if not structure:
        return
    
    for folder in structure:
        # Create indentation with proper spacing
        indent = "&nbsp;&nbsp;&nbsp;&nbsp;" * level
        
        # Choose emoji based on level
        if level == 0:
            emoji = "📁"
        elif level == 1:
            emoji = "📂"
        elif level == 2:
            emoji = "📄"
        else:
            emoji = "📋"
        
        # Create Google Drive link
        drive_link = folder.get('webViewLink', f"https://drive.google.com/drive/folders/{folder['id']}")
        
        # Display folder with proper indentation and clickable link
        folder_html = f"{indent}{emoji} <a href='{drive_link}' target='_blank'><strong>{folder['name']}</strong></a>"
        st.markdown(folder_html, unsafe_allow_html=True)
        
        # Recursively display children (show up to 4 levels for full hierarchy)
        if folder['children'] and level < 4:
            display_folder_structure(folder['children'], level + 1)

def filter_out_folders(structure, exclude_ids):
    """Recursively filter out folders with specified IDs"""
    if not structure:
        return []
    
    filtered_structure = []
    for folder in structure:
        # Skip if this folder should be excluded
        if folder['id'] in exclude_ids:
            continue
        
        # Recursively filter children
        filtered_children = filter_out_folders(folder['children'], exclude_ids)
        
        # Create new folder info with filtered children
        filtered_folder = {
            'id': folder['id'],
            'name': folder['name'],
            'level': folder['level'],
            'children': filtered_children
        }
        
        filtered_structure.append(filtered_folder)
    
    return filtered_structure

def upload_file_to_google_drive(service, file_data: bytes, file_name: str, mime_type: str, folder_id: str = None, shared_drive_id: str = None):
    """Upload a file to Google Drive (shared drive or personal)"""
    try:
        file_metadata = {
            'name': file_name
        }
        
        if folder_id:
            file_metadata['parents'] = [folder_id]
        
        # Create file stream
        file_stream = io.BytesIO(file_data)
        
        # Create MediaUpload object
        from googleapiclient.http import MediaIoBaseUpload
        media = MediaIoBaseUpload(
            file_stream,
            mimetype=mime_type,
            resumable=True
        )
        
        # Upload file with shared drive support if needed
        if shared_drive_id:
            file = service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, name, webViewLink',
                supportsAllDrives=True,
                supportsTeamDrives=True
            ).execute()
        else:
            file = service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, name, webViewLink'
            ).execute()
        
        return file
    except Exception as e:
        st.error(f"Error uploading file: {e}")
        return None

def check_existing_project_folder(service, molecule_code: str, parent_folder_id: str = None):
    """Check if a project folder with the same molecule code already exists"""
    try:
        project_folder_name = f"Project; Molecule {molecule_code}"
        
        if parent_folder_id:
            query = f"'{parent_folder_id}' in parents and name = '{project_folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        else:
            query = f"'{parent_folder_id or '0ALsvNdCE73XrUk9PVA'}' in parents and name = '{project_folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        
        # Use shared drive parameters for all queries
        results = service.files().list(
            q=query, 
            fields="files(id, name)",
            supportsAllDrives=True,
            supportsTeamDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        existing_folders = results.get('files', [])
        
        return existing_folders[0] if existing_folders else None
    except Exception as e:
        st.error(f"Error checking existing project folder: {e}")
        return None

def find_target_folder(service, molecule_code: str, campaign_number: str = None):
    """Find the Draft AI Reg Document -> IND -> Draft folder for the specified project"""
    try:
        # Find the project folder - use shared drive context
        project_folder = check_existing_project_folder(service, molecule_code, '0ALsvNdCE73XrUk9PVA')
        if not project_folder:
            st.error(f"❌ Project folder for Molecule {molecule_code} not found!")
            return None, None
        
        # Find Draft AI Reg Document folder
        reg_doc_query = f"'{project_folder['id']}' in parents and name = 'Draft AI Reg Document' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        reg_doc_results = service.files().list(
            q=reg_doc_query, 
            fields="files(id, name)",
            supportsAllDrives=True,
            supportsTeamDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        reg_doc_folders = reg_doc_results.get('files', [])
        
        if not reg_doc_folders:
            st.error(f"❌ Draft AI Reg Document folder not found in project {molecule_code}!")
            return None, None
        
        reg_doc_folder = reg_doc_folders[0]
        
        # Find IND folder
        ind_query = f"'{reg_doc_folder['id']}' in parents and name = 'IND' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        ind_results = service.files().list(
            q=ind_query, 
            fields="files(id, name)",
            supportsAllDrives=True,
            supportsTeamDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        ind_folders = ind_results.get('files', [])
        
        if not ind_folders:
            st.error(f"❌ IND folder not found in Draft AI Reg Document!")
            return None, None
        
        ind_folder = ind_folders[0]
        
        # Find Draft folder
        draft_query = f"'{ind_folder['id']}' in parents and name = 'Draft' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        draft_results = service.files().list(
            q=draft_query, 
            fields="files(id, name)",
            supportsAllDrives=True,
            supportsTeamDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        draft_folders = draft_results.get('files', [])
        
        if not draft_folders:
            st.error(f"❌ Draft folder not found in IND!")
            return None, None
        
        draft_folder = draft_folders[0]
        
        # Build the full path for display
        full_path = f"Project; Molecule {molecule_code} → Draft AI Reg Document → IND → Draft"
        
        return draft_folder, full_path
        
    except Exception as e:
        st.error(f"Error finding target folder: {e}")
        return None, None

def create_campaign_folder_structure(service, campaign_name: str, molecule_code: str, parent_folder_id: str = None, shared_drive_id: str = None):
    """Create the complete campaign folder structure as shown in the image"""
    try:
        # Set default parent folder ID if not provided
        if not parent_folder_id:
            parent_folder_id = '0ALsvNdCE73XrUk9PVA'  # New shared drive root folder
        
        # Check if project folder already exists
        existing_project = check_existing_project_folder(service, molecule_code, parent_folder_id)
        if existing_project:
            st.warning(f"⚠️ Project folder '{existing_project['name']}' already exists!")
            return None
        
        # Create main project folder
        project_folder_name = f"Project; Molecule {molecule_code}"
        project_folder = create_google_drive_folder(service, project_folder_name, parent_folder_id, shared_drive_id)
        if not project_folder:
            return None
        
        # Create campaign folder
        campaign_folder_name = f"Project {molecule_code} (Campaign #{campaign_name})"
        campaign_folder = create_google_drive_folder(service, campaign_folder_name, project_folder['id'], shared_drive_id)
        if not campaign_folder:
            return None
        
        # Create Draft AI Reg Document folder
        reg_doc_folder = create_google_drive_folder(service, "Draft AI Reg Document", project_folder['id'], shared_drive_id)
        if not reg_doc_folder:
            return None
        
        # Create Pre and Post folders under campaign
        pre_folder = create_google_drive_folder(service, "Pre", campaign_folder['id'], shared_drive_id)
        post_folder = create_google_drive_folder(service, "Post", campaign_folder['id'], shared_drive_id)
        
        # Create department folders under Pre and Post
        departments = ["mfg", "Anal", "Stability", "CTM"]
        statuses = ["Draft", "Review", "Approved"]
        
        for phase_folder in [pre_folder, post_folder]:
            if phase_folder:
                for dept in departments:
                    dept_folder = create_google_drive_folder(service, dept, phase_folder['id'], shared_drive_id)
                    if dept_folder:
                        for status in statuses:
                            create_google_drive_folder(service, status, dept_folder['id'], shared_drive_id)
        
        # Create regulatory document folders
        reg_types = ["IND", "IMPD", "Canada"]
        for reg_type in reg_types:
            reg_type_folder = create_google_drive_folder(service, reg_type, reg_doc_folder['id'], shared_drive_id)
            if reg_type_folder:
                for status in statuses:
                    create_google_drive_folder(service, status, reg_type_folder['id'], shared_drive_id)
        
        return {
            'project_folder': project_folder,
            'campaign_folder': campaign_folder,
            'reg_doc_folder': reg_doc_folder
        }
        
    except Exception as e:
        st.error(f"Error creating campaign folder structure: {e}")
        return None

def initialize_openai():
    """Initialize OpenAI client"""
    api_key = load_openai_api_key()
    if api_key and api_key != "your-openai-api-key-here":
        try:
            client = OpenAI(api_key=api_key)
            return client
        except Exception as e:
            return None
    return None

def create_sample_pharma_data():
    """Create sample pharmaceutical composition data"""
    data = {
        'Component': [
            'Active Pharmaceutical Ingredient',
            'Microcrystalline Cellulose',
            'Lactose Monohydrate',
            'Croscarmellose Sodium',
            'Magnesium Stearate',
            'Opadry II White'
        ],
        'Quality_Reference': ['USP', 'NF', 'NF', 'NF', 'NF', 'NF'],
        'Function': [
            'Active Ingredient',
            'Tablet Diluent',
            'Tablet Diluent',
            'Disintegrant',
            'Lubricant',
            'Film Coating'
        ],
        'Quantity_mg_per_tablet': [25.0, 150.0, 100.0, 10.0, 2.0, 8.0]
    }
    return pd.DataFrame(data)

def create_charts(df):
    """Create both Plotly and Matplotlib charts. Return (plotly_charts, export_charts) dicts."""
    import plotly.express as px
    plotly_charts = {}
    export_charts = {}
    # Plotly: Bar chart - Component quantities
    fig_bar = px.bar(df, x='Component', y='Quantity_mg_per_tablet', 
                     title='Component Quantities per Tablet',
                     color='Function', color_discrete_sequence=px.colors.qualitative.Set3)
    plotly_charts['component_quantities'] = fig_bar
    # Matplotlib: Bar chart
    fig, ax = plt.subplots(figsize=(6,4))
    df.groupby('Component')['Quantity_mg_per_tablet'].sum().plot(kind='bar', ax=ax, color='skyblue')
    ax.set_title('Component Quantities per Tablet')
    ax.set_ylabel('Quantity (mg/tablet)')
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png')
    buf.seek(0)
    export_charts['component_quantities'] = buf.read()
    plt.close(fig)
    # Plotly: Pie chart - Function distribution
    function_summary = df.groupby('Function')['Quantity_mg_per_tablet'].sum().reset_index()
    fig_pie = px.pie(function_summary, values='Quantity_mg_per_tablet', names='Function',
                     title='Distribution by Function', color_discrete_sequence=px.colors.qualitative.Set3)
    plotly_charts['function_distribution'] = fig_pie
    # Matplotlib: Pie chart
    fig, ax = plt.subplots(figsize=(6,4))
    function_summary2 = df.groupby('Function')['Quantity_mg_per_tablet'].sum()
    ax.pie(function_summary2, labels=function_summary2.index, autopct='%1.1f%%', startangle=90)
    ax.set_title('Distribution by Function')
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png')
    buf.seek(0)
    export_charts['function_distribution'] = buf.read()
    plt.close(fig)
    # Plotly: Bar chart - Quality reference distribution
    quality_counts = df['Quality_Reference'].value_counts()
    fig_quality = px.bar(x=quality_counts.index, y=quality_counts.values,
                        title='Quality Reference Distribution',
                        labels={'x': 'Quality Reference', 'y': 'Count'},
                        color=quality_counts.index, color_discrete_sequence=px.colors.qualitative.Set3)
    plotly_charts['quality_references'] = fig_quality
    # Matplotlib: Bar chart
    fig, ax = plt.subplots(figsize=(6,4))
    quality_counts.plot(kind='bar', ax=ax, color='lightgreen')
    ax.set_title('Quality Reference Distribution')
    ax.set_ylabel('Count')
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png')
    buf.seek(0)
    export_charts['quality_references'] = buf.read()
    plt.close(fig)
    # Plotly: Scatter plot - Component weight vs function
    fig_scatter = px.scatter(df, x='Function', y='Quantity_mg_per_tablet', 
                            size='Quantity_mg_per_tablet', color='Quality_Reference',
                            title='Component Weight vs Function',
                            hover_data=['Component'], color_discrete_sequence=px.colors.qualitative.Set3)
    plotly_charts['weight_vs_function'] = fig_scatter
    # Matplotlib: Scatter plot
    fig, ax = plt.subplots(figsize=(6,4))
    for func in df['Function'].unique():
        subset = df[df['Function'] == func]
        ax.scatter(subset['Function'], subset['Quantity_mg_per_tablet'], label=func, s=50)
    ax.set_title('Component Weight vs Function')
    ax.set_ylabel('Quantity (mg/tablet)')
    ax.legend()
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png')
    buf.seek(0)
    export_charts['weight_vs_function'] = buf.read()
    plt.close(fig)
    return plotly_charts, export_charts

def save_chart_as_image(fig_or_bytes, chart_name):
    if isinstance(fig_or_bytes, bytes):
        return fig_or_bytes
    return None

def generate_regulatory_text_with_ai(product_code: str, dosage_form: str, 
                                   composition_data: pd.DataFrame, 
                                   mechanism_of_action: str,
                                   drug_info: Dict,
                                   additional_instructions: str = "") -> Dict[str, str]:
    """Generate regulatory text using OpenAI only. If OpenAI is not available, raise an error."""
    client = initialize_openai()
    if not client:
        raise RuntimeError("OpenAI API key not found or invalid. Please set your OpenAI API key.")
    try:
        # Prepare composition data for AI
        composition_text = ""
        total_weight = 0
        for _, row in composition_data.iterrows():
            component = row['Component']
            quality_ref = row['Quality_Reference']
            function = row['Function']
            quantity = row['Quantity_mg_per_tablet']
            total_weight += quantity
            composition_text += f"- {component} ({quality_ref}): {quantity} mg ({function})\n"
        composition_text += f"- Total Weight: {total_weight} mg"
        # Build prompt with additional instructions
        additional_prompt = ""
        if additional_instructions.strip():
            additional_prompt = f"\nAdditional Instructions: {additional_instructions}"
        prompt = f"""
You are a regulatory-writing assistant drafting Section 3.2.P.1 "Description and Composition of the Drug Product" for an IND that is currently in Phase II clinical trials. All content must be suitable for direct inclusion in an eCTD‐compliant Module 3 dossier.

Product Information:
- Product code: {product_code}
- Dosage form: {dosage_form}
- Drug class: {drug_info.get('class', 'Not specified')}
- Indication: {drug_info.get('indication', 'Not specified')}
- Mechanism of action: {mechanism_of_action}

Composition data:
{composition_text}{additional_prompt}

Required Output Structure:
1. 3.2.P.1.1 Description of the Dosage Form
   - One concise paragraph that identifies the dosage form and strength(s)
   - States the active‐ingredient concentration(s) clearly
   - Summarises the mechanism of action in one sentence
   - Write in the third person, scientific style; do not use marketing language

2. 3.2.P.1.2 Composition
   - Introductory sentence: "The qualitative and quantitative composition of the {product_code} is provided in Table 1."
   - Table 1 should be titled 'Composition of the {product_code}'
   - Include all components with their quality references, functions, and quantities
   - Do not include markdown or ASCII tables in the text. Only refer to the table by title or as Table 1.

3. 3.2.P.1.3 Pharmaceutical Development
   - Brief description of formulation development considerations
   - Reference to key excipient functions

4. 3.2.P.1.4 Manufacturing Process
   - Overview of the manufacturing process
   - Key process parameters and controls

Please provide the text in a structured format suitable for regulatory submission.
"""
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a pharmaceutical regulatory writing expert with deep knowledge of FDA and ICH guidelines for IND submissions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.3
        )
        ai_text = response.choices[0].message.content
        sections = parse_ai_response(ai_text, product_code)
        return sections
    except Exception as e:
        raise RuntimeError(f"OpenAI API request failed: {e}")

def parse_ai_response(ai_text: str, product_code: str) -> Dict[str, str]:
    """Parse AI response into structured sections"""
    sections = {
        'description': '',
        'composition_intro': '',
        'pharmaceutical_development': '',
        'manufacturing_process': '',
        'table_title': f'Composition of the {product_code}'
    }
    
    # Simple parsing - in production, you might want more sophisticated parsing
    lines = ai_text.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        
        # Detect section headers
        if '3.2.P.1.1' in line or 'Description' in line:
            current_section = 'description'
            continue
        elif '3.2.P.1.2' in line or 'Composition' in line:
            current_section = 'composition_intro'
            continue
        elif '3.2.P.1.3' in line or 'Pharmaceutical Development' in line:
            current_section = 'pharmaceutical_development'
            continue
        elif '3.2.P.1.4' in line or 'Manufacturing Process' in line:
            current_section = 'manufacturing_process'
            continue
        elif line.startswith('Table') or line.startswith('The qualitative'):
            current_section = 'composition_intro'
        
        # Add content to appropriate section
        if current_section and line:
            if current_section == 'description':
                sections['description'] += line + ' '
            elif current_section == 'composition_intro':
                sections['composition_intro'] += line + ' '
            elif current_section == 'pharmaceutical_development':
                sections['pharmaceutical_development'] += line + ' '
            elif current_section == 'manufacturing_process':
                sections['manufacturing_process'] += line + ' '
    
    return sections

def export_to_word_regulatory(df: pd.DataFrame, sections: Dict[str, str], 
                             product_code: str, dosage_form: str,
                             uploaded_images: List, notes: str,
                             charts: Dict, chart_selections: Dict[str, bool]) -> Document:
    """Export regulatory document to Word format"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Section 3.2.P.1 Description and Composition of the Drug Product', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add description section
    doc.add_heading('3.2.P.1.1 Description of the Dosage Form', level=1)
    doc.add_paragraph(sections['description'])
    doc.add_paragraph()  # Add space
    
    # Add composition section
    doc.add_heading('3.2.P.1.2 Composition', level=1)
    doc.add_paragraph(sections['composition_intro'])
    doc.add_paragraph()  # Add space
    
    # Add table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    headers = ['Component', 'Quality Reference', 'Function', 'Quantity / Unit (mg per tablet)']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Add data rows
    total_weight = 0
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Component'])
        row_cells[1].text = str(row['Quality_Reference'])
        row_cells[2].text = str(row['Function'])
        row_cells[3].text = str(row['Quantity_mg_per_tablet'])
        total_weight += row['Quantity_mg_per_tablet']
    
    # Add total weight row
    total_cells = table.add_row().cells
    total_cells[0].text = 'Total Weight'
    total_cells[1].text = ''
    total_cells[2].text = ''
    total_cells[3].text = f'{total_weight}'
    
    # Add footnote
    doc.add_paragraph("Abbreviations: NF = National Formulary; Ph. Eur. = European Pharmacopoeia; USP = United States Pharmacopoeia.")
    
    # Add pharmaceutical development
    if sections.get('pharmaceutical_development'):
        doc.add_heading('3.2.P.1.3 Pharmaceutical Development', level=1)
        doc.add_paragraph(sections['pharmaceutical_development'])
        doc.add_paragraph()
    
    # Add manufacturing process
    if sections.get('manufacturing_process'):
        doc.add_heading('3.2.P.1.4 Manufacturing Process', level=1)
        doc.add_paragraph(sections['manufacturing_process'])
        doc.add_paragraph()
    
    # Add charts as images based on selection
    chart_selections = {
        'component_quantities': getattr(st.session_state, 'include_component_quantities', True),
        'function_distribution': getattr(st.session_state, 'include_function_distribution', True),
        'quality_references': getattr(st.session_state, 'include_quality_references', True),
        'weight_vs_function': getattr(st.session_state, 'include_weight_vs_function', True)
    }
    
    included_charts = {name: fig for name, fig in charts.items() if chart_selections.get(name, True)}
    
    if included_charts:
        doc.add_heading('Data Visualizations', level=1)
        for chart_name, fig in included_charts.items():
            # Try to save chart as image
            img_bytes = save_chart_as_image(fig, chart_name)
            
            if img_bytes:
                # Add image to document
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
                doc.add_paragraph(f"Figure: {chart_name.replace('_', ' ').title()}")
                doc.add_paragraph()
            else:
                # Fallback: add text description
                doc.add_paragraph(f"Chart: {chart_name.replace('_', ' ').title()}")
                doc.add_paragraph("(Chart visualization available in interactive version)")
                doc.add_paragraph()
    
    # Add notes if provided
    if notes.strip():
        doc.add_heading('Additional Notes', level=1)
        doc.add_paragraph(notes)
    
    # Add images if provided
    if uploaded_images:
        doc.add_heading('Supporting Images', level=1)
        for i, img_data in enumerate(uploaded_images):
            try:
                img_stream = io.BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(4))
                doc.add_paragraph(f"Figure {i+1}")
            except Exception as e:
                doc.add_paragraph(f"Image {i+1} could not be added: {str(e)}")
    
    return doc

def export_to_pdf_regulatory(df: pd.DataFrame, sections: Dict[str, str], 
                            product_code: str, dosage_form: str,
                            uploaded_images: List, notes: str,
                            charts: Dict, chart_selections: Dict[str, bool]):
    """Export regulatory document to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1
    )
    story.append(Paragraph('Section 3.2.P.1 Description and Composition of the Drug Product', title_style))
    story.append(Spacer(1, 12))
    
    # Add description section
    story.append(Paragraph('3.2.P.1.1 Description of the Dosage Form', styles['Heading2']))
    story.append(Paragraph(sections['description'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Add composition section
    story.append(Paragraph('3.2.P.1.2 Composition', styles['Heading2']))
    story.append(Paragraph(sections['composition_intro'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Prepare table data
    headers = ['Component', 'Quality Reference', 'Function', 'Quantity / Unit (mg per tablet)']
    table_data = [headers]
    
    total_weight = 0
    for _, row in df.iterrows():
        table_data.append([
            str(row['Component']),
            str(row['Quality_Reference']),
            str(row['Function']),
            str(row['Quantity_mg_per_tablet'])
        ])
        total_weight += row['Quantity_mg_per_tablet']
    
    # Add total weight row
    table_data.append(['Total Weight', '', '', str(total_weight)])
    
    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    story.append(table)
    story.append(Spacer(1, 12))
    
    # Add footnote
    story.append(Paragraph("Abbreviations: NF = National Formulary; Ph. Eur. = European Pharmacopoeia; USP = United States Pharmacopoeia.", styles['Normal']))
    
    # Add pharmaceutical development
    if sections.get('pharmaceutical_development'):
        story.append(Paragraph('3.2.P.1.3 Pharmaceutical Development', styles['Heading2']))
        story.append(Paragraph(sections['pharmaceutical_development'], styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add manufacturing process
    if sections.get('manufacturing_process'):
        story.append(Paragraph('3.2.P.1.4 Manufacturing Process', styles['Heading2']))
        story.append(Paragraph(sections['manufacturing_process'], styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add charts as images based on selection
    chart_selections = {
        'component_quantities': getattr(st.session_state, 'include_component_quantities', True),
        'function_distribution': getattr(st.session_state, 'include_function_distribution', True),
        'quality_references': getattr(st.session_state, 'include_quality_references', True),
        'weight_vs_function': getattr(st.session_state, 'include_weight_vs_function', True)
    }
    
    included_charts = {name: fig for name, fig in charts.items() if chart_selections.get(name, True)}
    
    if included_charts:
        story.append(Paragraph('Data Visualizations', styles['Heading2']))
        for chart_name, fig in included_charts.items():
            # Try to save chart as image
            img_bytes = save_chart_as_image(fig, chart_name)
            
            if img_bytes:
                # Add image to PDF
                img_stream = io.BytesIO(img_bytes)
                img = RLImage(img_stream, width=5*inch, height=3*inch)
                story.append(img)
                story.append(Paragraph(f"Figure: {chart_name.replace('_', ' ').title()}", styles['Normal']))
                story.append(Spacer(1, 12))
            else:
                # Fallback: add text description
                story.append(Paragraph(f"Chart: {chart_name.replace('_', ' ').title()}", styles['Normal']))
                story.append(Paragraph("(Chart visualization available in interactive version)", styles['Normal']))
                story.append(Spacer(1, 12))
    
    # Add notes if provided
    if notes.strip():
        story.append(Paragraph('Additional Notes', styles['Heading2']))
        story.append(Paragraph(notes, styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add images if provided
    if uploaded_images:
        story.append(Paragraph('Supporting Images', styles['Heading2']))
        for i, img_data in enumerate(uploaded_images):
            try:
                img_stream = io.BytesIO(img_data)
                img = RLImage(img_stream, width=3*inch, height=2*inch)
                story.append(img)
                story.append(Paragraph(f"Figure {i+1}", styles['Normal']))
                story.append(Spacer(1, 12))
            except Exception as e:
                story.append(Paragraph(f"Image {i+1} could not be added: {str(e)}", styles['Normal']))
    
    doc.build(story)
    return buffer

def main():
    # Header
    st.markdown('<h1 class="main-header">📁 Document Generator and Folder Management System</h1>', unsafe_allow_html=True)
    st.markdown("Generate regulatory documents and manage Google Drive folder structures for pharmaceutical projects")
    
    # Initialize services
    drive_service = initialize_google_drive_service()
    
    # Get shared drive ID if available
    shared_drive_id = None
    if drive_service:
        # Display service account email for shared drive setup
        creds = load_google_drive_credentials()
        
        shared_drive_id = get_shared_drive_id(drive_service)
        st.session_state.shared_drive_id = shared_drive_id
    
    # Top configuration section
    st.markdown('<h2 class="section-header">⚙️ Configuration</h2>', unsafe_allow_html=True)
    
    st.subheader("📋 Product Configuration")
    # Drug selection
    selected_drug = st.selectbox("Select Drug", list(DRUG_DATABASE.keys()), key="selected_drug")
    drug_info = DRUG_DATABASE[selected_drug]
    
    # Display drug information
    st.markdown('<div class="drug-info">', unsafe_allow_html=True)
    st.write(f"**Drug Class:** {drug_info['class']}")
    st.write(f"**Indication:** {drug_info['indication']}")
    st.write(f"**Available Strengths:** {drug_info['strength']}")
    st.write(f"**Manufacturer:** {drug_info['manufacturer']}")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Product information
    product_code = st.text_input("Product Code", f"{selected_drug}-001", key="product_code")
    dosage_form = st.text_input("Dosage Form", drug_info['dosage_form'], key="dosage_form")
    mechanism_of_action = st.text_area("Mechanism of Action", drug_info['mechanism'], height=100, key="mechanism_of_action")
    
    # System Status Section
    st.markdown('<h2 class="section-header">🔧 System Status</h2>', unsafe_allow_html=True)
    
    # Check OpenAI status
    openai_client = initialize_openai()
    if openai_client:
        st.success("✅ OpenAI API connected")
    else:
        st.error("❌ OpenAI API not connected")
        st.info("ℹ️ Add 'OPENAI_API_KEY' to credentials.py or set it in Streamlit secrets")
    
    # Check Google Drive status
    if drive_service:
        st.success("✅ Google Drive API connected")
    else:
        st.error("❌ Google Drive API not connected")
        st.info("ℹ️ Add 'google_drive_api' to Streamlit secrets or place 'aaitdemoharmony-3945571299f1.json' in the app directory")
    
    if st.session_state.shared_drive_id:
        st.info(f"📁 Using Shared Drive ID: {st.session_state.shared_drive_id}")
    else:
        st.warning("⚠️ No shared drive found. File uploads may fail due to service account storage limitations.")
        st.info("💡 Create a shared drive and add the service account email above as a member with 'Editor' permissions") 
    
    # Display current folder structure if connected
    if drive_service:
        st.markdown('<h2 class="section-header">📂 Current Google Drive Structure</h2>', unsafe_allow_html=True)
        
        #show_debug = st.checkbox("Show debug info", key="show_debug")
        
        if st.button("📁 Load Structure", key="load_structure"):
            with st.spinner("Loading folder structure..."):
                try:
                    # Pass the shared drive ID to ensure proper folder listing
                    # Increase max_depth to 5 to get full hierarchy
                    structure = get_folder_structure_recursive(drive_service, st.session_state.shared_drive_id, max_depth=5)
                    
                    if structure:
                        st.session_state.folder_structure = structure
                        st.success(f"✅ Loaded {len(structure)} folders")
                        

                    else:
                        st.info("📁 No folders found")
                        st.session_state.folder_structure = None
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                    st.session_state.folder_structure = None
        
        if 'folder_structure' in st.session_state and st.session_state.folder_structure:
            st.subheader("📂 Folder Structure")
            with st.expander("View Folder Structure", expanded=True):
                display_folder_structure(st.session_state.folder_structure)
            

        else:
            st.info("Click 'Load Structure' to view your Google Drive folders")
        
        # Campaign Management Section
        st.markdown('<h2 class="section-header">🚀 Campaign Management</h2>', unsafe_allow_html=True)
        
        st.subheader("📋 Campaign Details")
        molecule_code = st.text_input("Molecule Code", "THPG001", key="molecule_code")
        campaign_number = st.text_input("Campaign Number", "3", key="campaign_number")
        
        st.subheader("🚀 Create Campaign")
        if st.button("🚀 Create Campaign Structure", type="primary", key="create_campaign"):
            if molecule_code and campaign_number:
                with st.spinner("Creating campaign folder structure..."):
                    try:
                        result = create_campaign_folder_structure(
                            drive_service, 
                            campaign_number, 
                            molecule_code,
                            None,  # parent_folder_id
                            shared_drive_id
                        )
                        if result:
                            st.success(f"✅ Campaign structure created successfully!")
                            st.write(f"**Project Folder:** {result['project_folder']['name']}")
                            st.write(f"**Campaign Folder:** {result['campaign_folder']['name']}")
                            st.write(f"**Reg Doc Folder:** {result['reg_doc_folder']['name']}")
                            
                            # Refresh the folder structure
                            st.rerun()
                        else:
                            st.error("❌ Failed to create campaign structure")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
            else:
                st.warning("Please enter both Molecule Code and Campaign Number")
    
    # Data input section
    st.markdown('<h2 class="section-header">📊 Composition Data</h2>', unsafe_allow_html=True)
    
    # Option to upload CSV or use sample data
    data_option = st.radio("Choose data source:", ["Use Sample Data", "Upload CSV File"], key="data_option")
    
    if data_option == "Upload CSV File":
        uploaded_file = st.file_uploader("Choose a CSV file", type="csv", key="csv_uploader")
        if uploaded_file is not None:
            df = pd.read_csv(uploaded_file)
            st.success(f"✅ Loaded {len(df)} components")
        else:
            st.info("Please upload a CSV file or use sample data")
            df = None
    else:
        df = create_sample_pharma_data()
        st.success("✅ Using sample pharmaceutical data")
    
    if df is not None:
        # Display the composition data
        st.subheader("📋 Composition Table")
        st.dataframe(df, use_container_width=True)
        
        # Calculate total weight
        total_weight = df['Quantity_mg_per_tablet'].sum()
        st.info(f"📊 Total tablet weight: {total_weight} mg")
        
        # Create and display charts
        st.markdown('<h2 class="section-header">📈 Data Visualizations</h2>', unsafe_allow_html=True)
        plotly_charts, export_charts = create_charts(df)
        
        # Display charts in tabs
        chart_tabs = st.tabs(["Component Quantities", "Function Distribution", "Quality References", "Weight vs Function"])
        
        with chart_tabs[0]:
            st.plotly_chart(plotly_charts['component_quantities'], use_container_width=True)
            if 'include_component_quantities' in locals() and include_component_quantities:
                st.success("✅ This chart will be included in the report")
            else:
                st.info("ℹ️ Use the chart selection options to include this chart in the report")
        
        with chart_tabs[1]:
            st.plotly_chart(plotly_charts['function_distribution'], use_container_width=True)
            if 'include_function_distribution' in locals() and include_function_distribution:
                st.success("✅ This chart will be included in the report")
            else:
                st.info("ℹ️ Use the chart selection options to include this chart in the report")
        
        with chart_tabs[2]:
            st.plotly_chart(plotly_charts['quality_references'], use_container_width=True)
            if 'include_quality_references' in locals() and include_quality_references:
                st.success("✅ This chart will be included in the report")
            else:
                st.info("ℹ️ Use the chart selection options to include this chart in the report")
        
        with chart_tabs[3]:
            st.plotly_chart(plotly_charts['weight_vs_function'], use_container_width=True)
            if 'include_weight_vs_function' in locals() and include_weight_vs_function:
                st.success("✅ This chart will be included in the report")
            else:
                st.info("ℹ️ Use the chart selection options to include this chart in the report")
        
        # Initialize variables for additional content
        uploaded_images = []
        notes = ""
        compliance_level = "FDA"
        document_version = "1.0"
        include_component_quantities = True
        include_function_distribution = True
        include_quality_references = True
        include_weight_vs_function = True
        
        # Additional content section (collapsible)
        with st.expander("📎 Additional Content", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📷 Upload Images")
                uploaded_images = st.file_uploader("Upload supporting images", 
                                                 type=['png', 'jpg', 'jpeg'], 
                                                 accept_multiple_files=True,
                                                 key="image_uploader")
                
                if uploaded_images:
                    st.write(f"✅ {len(uploaded_images)} image(s) uploaded")
                    for i, img in enumerate(uploaded_images):
                        st.image(img, caption=f"Image {i+1}: {img.name}", width=200)
            
            with col2:
                st.subheader("📝 Additional Notes")
                notes = st.text_area("Add any additional notes or comments for the document", 
                                    height=150,
                                    placeholder="Enter any additional information, observations, or notes...",
                                    key="additional_notes")
                
                st.subheader("🔧 Document Options")
                compliance_level = st.selectbox("Compliance level:", ["FDA", "EMA", "ICH", "Other"], key="compliance_level")
                document_version = st.text_input("Document version:", "1.0", key="document_version")
                
                st.subheader("📊 Chart Selection for Report")
                include_component_quantities = st.checkbox("Include Component Quantities Chart", value=True, key="include_component_quantities")
                include_function_distribution = st.checkbox("Include Function Distribution Chart", value=True, key="include_function_distribution")
                include_quality_references = st.checkbox("Include Quality References Chart", value=True, key="include_quality_references")
                include_weight_vs_function = st.checkbox("Include Weight vs Function Chart", value=True, key="include_weight_vs_function")
        
        # Generate regulatory text
        st.markdown('<h2 class="section-header">📝 Generated Regulatory Text</h2>', unsafe_allow_html=True)
        
        if st.button("🔄 Generate Regulatory Text", type="primary", key="generate_text"):
            with st.spinner("Generating regulatory text..."):
                if initialize_openai():
                    sections = generate_regulatory_text_with_ai(
                        product_code, dosage_form, df, mechanism_of_action, drug_info
                    )
                else:
                    raise RuntimeError("OpenAI API key not found or invalid. Please set your OpenAI API key.")
                # Store in session state for export
                st.session_state.sections = sections
                st.session_state.df = df
                st.session_state.uploaded_images = uploaded_images if uploaded_images else []
                st.session_state.notes = notes
                st.session_state.charts = export_charts
                st.session_state.text_generated = True

        # Always display the current text if it exists
        if st.session_state.get('sections'):
            st.markdown('<div class="regulatory-text">', unsafe_allow_html=True)
            st.subheader("3.2.P.1.1 Description of the Dosage Form")
            st.write(st.session_state.sections.get('description', ''))
            if st.session_state.sections.get('composition_intro') and st.session_state.sections['composition_intro'].strip():
                st.subheader("3.2.P.1.2 Composition")
                st.write(st.session_state.sections['composition_intro'])
            if st.session_state.sections.get('pharmaceutical_development') and st.session_state.sections['pharmaceutical_development'].strip():
                st.subheader("3.2.P.1.3 Pharmaceutical Development")
                st.write(st.session_state.sections['pharmaceutical_development'])
            if st.session_state.sections.get('manufacturing_process') and st.session_state.sections['manufacturing_process'].strip():
                st.subheader("3.2.P.1.4 Manufacturing Process")
                st.write(st.session_state.sections['manufacturing_process'])
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Export section
        if 'sections' in st.session_state:
            st.markdown('<h2 class="section-header">💾 Export Document</h2>', unsafe_allow_html=True)
            
            # Create chart selections dictionary (shared by all export functions)
            chart_selections = {
                'component_quantities': include_component_quantities,
                'function_distribution': include_function_distribution,
                'quality_references': include_quality_references,
                'weight_vs_function': include_weight_vs_function
            }
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.subheader("Export to Word (.docx)")
                if st.button("📄 Generate Word Document", type="primary", key="generate_word"):
                    with st.spinner("Generating Word document..."):
                        # Convert uploaded images to bytes
                        image_bytes = []
                        if st.session_state.uploaded_images:
                            for img in st.session_state.uploaded_images:
                                image_bytes.append(img.read())
                                img.seek(0)  # Reset file pointer
                        
                        doc = export_to_word_regulatory(
                            st.session_state.df, 
                            st.session_state.sections, 
                            product_code, 
                            dosage_form,
                            image_bytes,
                            st.session_state.notes,
                            st.session_state.charts,
                            chart_selections
                        )
                        
                        # Save to bytes
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        
                        # Create download button
                        st.download_button(
                            label="📥 Download Word Document",
                            data=doc_buffer.getvalue(),
                            file_name=f"Section_3.2.P.1_{product_code}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            
            with col2:
                st.subheader("Export to PDF")
                if st.button("📄 Generate PDF Document", type="primary", key="generate_pdf"):
                    with st.spinner("Generating PDF document..."):
                        # Convert uploaded images to bytes
                        image_bytes = []
                        if st.session_state.uploaded_images:
                            for img in st.session_state.uploaded_images:
                                image_bytes.append(img.read())
                                img.seek(0)  # Reset file pointer
                        
                        pdf_buffer = export_to_pdf_regulatory(
                            st.session_state.df, 
                            st.session_state.sections, 
                            product_code, 
                            dosage_form,
                            image_bytes,
                            st.session_state.notes,
                            st.session_state.charts,
                            chart_selections
                        )
                        pdf_buffer.seek(0)
                        
                        # Create download button
                        st.download_button(
                            label="📥 Download PDF Document",
                            data=pdf_buffer.getvalue(),
                            file_name=f"Section_3.2.P.1_{product_code}.pdf",
                            mime="application/pdf"
                        )
            
            with col3:
                st.subheader("📁 Upload to Google Drive")
                if drive_service:
                    # Get molecule code and campaign number for upload
                    upload_molecule_code = st.text_input("Molecule Code for upload:", molecule_code, key="upload_molecule_code")
                    upload_campaign_number = st.text_input("Campaign Number for upload:", campaign_number, key="upload_campaign_number")
                    
                    # Find target folder
                    target_folder, folder_path = find_target_folder(drive_service, upload_molecule_code, upload_campaign_number)
                    
                    if target_folder:
                        st.success(f"✅ Target folder found: {folder_path}")
                        st.info(f"📁 Folder ID: {target_folder['id']}")
                        st.info(f"🔗 Folder Link: https://drive.google.com/drive/folders/{target_folder['id']}")
                    else:
                        st.warning("⚠️ Target folder not found. Please ensure the project and campaign structure exists.")
                    
                    if st.button("☁️ Upload Word to Drive", type="primary", key="upload_word"):
                        if target_folder:
                            with st.spinner("Generating and uploading Word document..."):
                                try:
                                    # Generate document
                                    image_bytes = []
                                    if st.session_state.uploaded_images:
                                        for img in st.session_state.uploaded_images:
                                            image_bytes.append(img.read())
                                            img.seek(0)
                                    
                                    doc = export_to_word_regulatory(
                                        st.session_state.df, 
                                        st.session_state.sections, 
                                        product_code, 
                                        dosage_form,
                                        image_bytes,
                                        st.session_state.notes,
                                        st.session_state.charts,
                                        chart_selections
                                    )
                                    
                                    # Save to bytes
                                    doc_buffer = io.BytesIO()
                                    doc.save(doc_buffer)
                                    doc_buffer.seek(0)
                                    
                                    # Upload to Google Drive
                                    file_name = f"Section_3.2.P.1_{product_code}.docx"
                                    uploaded_file = upload_file_to_google_drive(
                                        drive_service,
                                        doc_buffer.getvalue(),
                                        file_name,
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        target_folder['id'],
                                        shared_drive_id
                                    )
                                    
                                    if uploaded_file:
                                        st.success(f"✅ Uploaded to Google Drive: {uploaded_file['name']}")
                                        st.write(f"**Project:** Molecule {upload_molecule_code}")
                                        st.write(f"**Campaign:** {upload_campaign_number}")
                                        st.write(f"**Upload Path:** {folder_path}")
                                        st.write(f"**File Link:** {uploaded_file.get('webViewLink', 'No link available')}")
                                        st.write(f"**Folder Link:** https://drive.google.com/drive/folders/{target_folder['id']}")
                                    else:
                                        st.error("❌ Failed to upload to Google Drive")
                                except Exception as e:
                                    st.error(f"❌ Error: {e}")
                        else:
                            st.warning("Please ensure the target folder is found before uploading")
                    
                    if st.button("☁️ Upload PDF to Drive", type="primary", key="upload_pdf"):
                        if target_folder:
                            with st.spinner("Generating and uploading PDF document..."):
                                try:
                                    # Generate document
                                    image_bytes = []
                                    if st.session_state.uploaded_images:
                                        for img in st.session_state.uploaded_images:
                                            image_bytes.append(img.read())
                                            img.seek(0)
                                    
                                    pdf_buffer = export_to_pdf_regulatory(
                                        st.session_state.df, 
                                        st.session_state.sections, 
                                        product_code, 
                                        dosage_form,
                                        image_bytes,
                                        st.session_state.notes,
                                        st.session_state.charts,
                                        chart_selections
                                    )
                                    pdf_buffer.seek(0)
                                    
                                    # Upload to Google Drive
                                    file_name = f"Section_3.2.P.1_{product_code}.pdf"
                                    uploaded_file = upload_file_to_google_drive(
                                        drive_service,
                                        pdf_buffer.getvalue(),
                                        file_name,
                                        "application/pdf",
                                        target_folder['id'],
                                        shared_drive_id
                                    )
                                    
                                    if uploaded_file:
                                        st.success(f"✅ Uploaded to Google Drive: {uploaded_file['name']}")
                                        st.write(f"**Project:** Molecule {upload_molecule_code}")
                                        st.write(f"**Campaign:** {upload_campaign_number}")
                                        st.write(f"**Upload Path:** {folder_path}")
                                        st.write(f"**File Link:** {uploaded_file.get('webViewLink', 'No link available')}")
                                        st.write(f"**Folder Link:** https://drive.google.com/drive/folders/{target_folder['id']}")
                                    else:
                                        st.error("❌ Failed to upload to Google Drive")
                                except Exception as e:
                                    st.error(f"❌ Error: {e}")
                        else:
                            st.warning("Please ensure the target folder is found before uploading")
                else:
                    st.info("Connect to Google Drive to upload documents")

if __name__ == "__main__":
    main()
